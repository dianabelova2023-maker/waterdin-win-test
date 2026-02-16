from __future__ import annotations

from io import BytesIO
from typing import Dict, List

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(0)


def _fmt(v: float, digits: int = 2) -> str:
    return f"{float(v or 0.0):.{digits}f}".replace(".", ",")


def _add_row_merged(table, idx: str, label: str, value: str):
    r = table.add_row().cells
    r[0].text = idx
    r[1].text = label
    r[2].text = ""
    r[3].text = value
    r[1].merge(r[2])
    return r


def _set_col_widths(table, widths_cm: List[float]) -> None:
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _disable_row_split(table) -> None:
    # Запрещает разрыв строк таблицы между страницами (как в образце формы).
    for row in table.rows:
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def build_gvs_passport_docx(
    object_name: str,
    water_rows: List[Dict[str, float | str]],
    gvs: Dict[str, float],
    passport_inputs: Dict[str, float | str | dict],
    annex_label: str = "",
) -> bytes:
    doc = Document()
    _set_doc_defaults(doc)

    annex_text = (annex_label or "").strip()
    p_annex = doc.add_paragraph()
    p_annex.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_annex.add_run(f"Приложение {annex_text or '____'}")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Паспорт ГВС").bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "№\nп/п"
    h[1].text = "Наименование показателя"
    h[2].text = ""
    h[3].text = "Данные по проекту"
    h[1].merge(h[2])
    _set_col_widths(table, [1.2, 8.2, 5.2, 4.0])

    total_count = float(passport_inputs.get("consumers_count", 0.0) or 0.0)
    t_hours = float(passport_inputs.get("hours_per_day", 12.0) or 12.0)
    area_m2 = float(passport_inputs.get("area_m2", 0.0) or 0.0)
    volume_m3 = float(passport_inputs.get("volume_m3", 0.0) or 0.0)
    devices_total = int(float(passport_inputs.get("devices_total", 0.0) or 0.0))
    fixture_name = str(passport_inputs.get("fixture_name", "Раковина"))
    q0_char = float(passport_inputs.get("q0_char_l_s", 0.14) or 0.14)
    q0hr_char = float(passport_inputs.get("q0hr_char_l_h", 60.0) or 60.0)
    h_top = float(passport_inputs.get("h_top_m", 0.0) or 0.0)
    free_head_m = float(passport_inputs.get("free_head_m", 20.0) or 20.0)
    losses_system_m = float(passport_inputs.get("losses_system_m", 0.0) or 0.0)
    circ_losses_m = float(passport_inputs.get("circ_losses_m", 0.0) or 0.0)
    has_meter = bool(passport_inputs.get("has_meter", True))
    has_itp_heating = bool(passport_inputs.get("has_itp_heating", False))
    meter_loss_m = float(passport_inputs.get("meter_loss_m", 0.0) or 0.0)
    fixture_counts = passport_inputs.get("fixture_counts", {}) or {}

    q_day_hot = float(gvs.get("qh_avg_m3_h", 0.0) or 0.0) * t_hours
    q_avg_hour = float(gvs.get("qh_avg_m3_h", 0.0) or 0.0)
    q_max_hour = float(gvs.get("qh_max_m3_h", 0.0) or 0.0)
    q_sec = float(gvs.get("qh_l_s", 0.0) or 0.0)

    # Графа 8: вероятность действия приборов, по СП 30 (формула P).
    # Берем q_hr,u по ГВС как средневзвешенный на 1 потребителя из расчетных строк воды.
    p_action = 0.0
    sum_u = 0.0
    sum_qhr_hot_u = 0.0
    for r in water_rows:
        u = float(r.get("count", 0.0) or 0.0)
        if u <= 0.0:
            continue
        qhr_hot_u = float(r.get("q_hr_hot_l_h", 0.0) or 0.0)
        if qhr_hot_u <= 0.0:
            continue
        sum_u += u
        sum_qhr_hot_u += qhr_hot_u * u
    qhru_hot_l_h = (sum_qhr_hot_u / sum_u) if sum_u > 0 else 0.0
    if devices_total > 0 and q0_char > 0 and qhru_hot_l_h > 0:
        p_action = qhru_hot_l_h * max(total_count, 0.0) / (3600.0 * q0_char * devices_total)

    # Графа 9: вероятность использования приборов, по СП 30 (формула Phr = 3600*P*q0/q0hr).
    p_use = 0.0
    if q0hr_char > 0 and p_action > 0:
        p_use = 3600.0 * p_action * q0_char / q0hr_char

    q_heat_avg_kcal_h = float(gvs.get("qth_kW", 0.0) or 0.0) * 860.0
    q_heat_max_kcal_h = float(gvs.get("qhrh_kW", 0.0) or 0.0) * 860.0
    q_heat_spec_kcal_h_m2 = q_heat_avg_kcal_h / area_m2 if area_m2 > 0 else 0.0
    # Строка 18: только потери в системе (без свободного напора).
    # Строка 19: включает свободный напор.
    pressure_losses_m = losses_system_m
    h_hex_m = 3.0 if has_itp_heating else 0.0
    pressure_required_m = h_top + pressure_losses_m + free_head_m + (meter_loss_m if has_meter else 0.0) + h_hex_m
    qht_kcal_h = float(gvs.get("qht_kW", 0.0) or 0.0) * 860.0
    qcir_l_s = float(gvs.get("qcir_l_s", 0.0) or 0.0)

    _add_row_merged(table, "1", "Назначение здания", object_name)
    _add_row_merged(table, "2", "Количество основных потребителей", _fmt(total_count, 0))
    _add_row_merged(table, "3", "Общая площадь, м²", _fmt(area_m2, 1) if area_m2 > 0 else "-")
    _add_row_merged(table, "4", "Строительный объем, м³", _fmt(volume_m3, 1) if volume_m3 > 0 else "-")
    _add_row_merged(table, "5", "Общее количество санитарных приборов, шт.", _fmt(devices_total, 0))

    fixture_rows = []
    for fname, fcnt in fixture_counts.items():
        if float(fcnt or 0.0) > 0:
            r = table.add_row().cells
            r[0].text = ""
            r[1].text = ""
            r[2].text = str(fname)
            r[3].text = _fmt(float(fcnt), 0)
            fixture_rows.append(r)

    if fixture_rows:
        fixture_rows[0][1].text = ""
        fixture_rows[0][1].merge(fixture_rows[-1][1])
        fixture_rows[0][0].merge(fixture_rows[-1][0])

    _add_row_merged(table, "6", "Число часов работы в сутки, ч/сут", _fmt(t_hours, 0))
    _add_row_merged(
        table,
        "7",
        f"Расход воды характерным прибором ({fixture_name}, Приложение А.1 СП 30.13330.2020), л/с (л/ч)",
        f"{_fmt(q0_char, 3)} ({_fmt(q0hr_char, 0)})",
    )
    _add_row_merged(table, "8", "Вероятность действия водоразборных приборов", _fmt(p_action, 4))
    _add_row_merged(table, "9", "Вероятность использования водоразборных приборов", _fmt(p_use, 4))

    # 10-13: блок "Расчетные расходы воды"
    w_rows = []
    for idx, sub, val in [
        ("10", "Секундный, л/с", _fmt(q_sec, 2)),
        ("11", "Суточный, м³/сут", _fmt(q_day_hot, 2)),
        ("12", "Средний часовой, м³/ч", _fmt(q_avg_hour, 2)),
        ("13", "Максимальный часовой, м³/ч", _fmt(q_max_hour, 2)),
    ]:
        r = table.add_row().cells
        r[0].text = idx
        r[1].text = ""
        r[2].text = sub
        r[3].text = val
        w_rows.append(r)
    w_merged = w_rows[0][1].merge(w_rows[-1][1])
    w_merged.text = "Расчетные расходы воды"

    # 14-16: блок "Расход тепла"
    h_rows = []
    for idx, sub, val in [
        ("14", "Средний часовой, ккал/ч", _fmt(q_heat_avg_kcal_h, 0)),
        ("15", "Максимальный часовой, ккал/ч", _fmt(q_heat_max_kcal_h, 0)),
        ("16", "Удельный (на 1 м² площади), ккал/ч·м²", _fmt(q_heat_spec_kcal_h_m2, 1) if q_heat_spec_kcal_h_m2 > 0 else "-"),
    ]:
        r = table.add_row().cells
        r[0].text = idx
        r[1].text = ""
        r[2].text = sub
        r[3].text = val
        h_rows.append(r)
    h_merged = h_rows[0][1].merge(h_rows[-1][1])
    h_merged.text = "Расход тепла"

    _add_row_merged(table, "17", "Высота диктующего прибора (трубопровода) над точкой подключения, м", _fmt(h_top, 2))
    _add_row_merged(table, "18", "Потери давления в системе, м вод. ст.", _fmt(pressure_losses_m, 2))
    row19_label = "Необходимое давление на выходе из ТП/ИТП, включая свободный напор, м"
    row19_val = _fmt(pressure_required_m, 2)
    _add_row_merged(table, "19", row19_label, row19_val)
    _add_row_merged(table, "20", "Потери тепла трубопроводами, ккал/ч", _fmt(qht_kcal_h, 0))
    _add_row_merged(table, "21", "Расход воды на циркуляцию, л/с", _fmt(qcir_l_s, 3))
    _add_row_merged(table, "22", "Потери давления в циркуляционном кольце, м вод. ст.", _fmt(circ_losses_m, 2))
    _add_row_merged(table, "23", "Температура горячей воды, °C", _fmt(float(gvs.get("t_hot_c", 0.0) or 0.0), 1))

    for row in table.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Центрирование крупных объединенных подписей блоков.
    if w_rows:
        w_rows[0][1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h_rows:
        h_rows[0][1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fixture_rows:
        fixture_rows[0][1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    _disable_row_split(table)

    note = doc.add_paragraph()
    note.add_run(
        "Расчет выполнен по СП 30.13330.2020 (в т.ч. формулы 12, 13, 16, 17) "
        "и нормативным значениям Приложений А.1 и А.2."
    ).italic = True

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
