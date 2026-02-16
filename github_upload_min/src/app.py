from __future__ import annotations

import csv
import os
import sys
import base64
import re
from datetime import datetime
from io import BytesIO, StringIO
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont

from calcs import (
    WaterConsumer,
    build_data_checks,
    calc_gvs_passport,
    calc_water_by_consumers_advanced,
)
from hydraulics import (
    CAST_IRON_BY_CLASS,
    CAST_IRON_DIMENSIONS,
    COPPER_DIMENSIONS,
    FIBERGLASS_DIMENSIONS,
    K_PRESETS,
    MLPEX_SDR_SERIES,
    MATERIALS,
    METAL_PLASTIC_DIMENSIONS,
    METAL_PLASTIC_ID_MM,
    PLASTIC_PE_GRADES,
    PLASTIC_DIMENSIONS,
    POLYPLASTIC_DIMENSIONS,
    PLASTIC_SDR_SERIES,
    POLYPLASTIC_ID_MM,
    STEEL_DIMENSIONS,
    calc_hydraulics,
)
from passport_gvs_docx import build_gvs_passport_docx
from report_docx import build_report_docx


st.set_page_config(page_title="Waterdin", page_icon="💧", layout="wide")

IS_NATIVE_APP = (os.getenv("WATERDIN_NATIVE", "0") == "1") or bool(getattr(sys, "frozen", False))

ICON_PATH = Path(__file__).resolve().parents[1] / "assets" / "waterdin_icon.png"
BG_ICON_BASE64 = ""
try:
    BG_ICON_BASE64 = base64.b64encode(ICON_PATH.read_bytes()).decode("ascii")
except Exception:
    BG_ICON_BASE64 = ""

css_theme = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Manrope:wght@500;700;800&family=Pacifico&display=swap');

html, body, [class*="css"]  {
    font-family: "Manrope", sans-serif;
}

.din-script {
    font-family: "Pacifico", cursive;
    font-weight: 400;
    letter-spacing: 0.4px;
}

[data-testid="stAppViewContainer"] {
    background:
      radial-gradient(1200px 600px at 8% -10%, rgba(179, 223, 255, 0.45), transparent 60%),
      radial-gradient(1000px 500px at 95% -15%, rgba(150, 210, 255, 0.40), transparent 55%),
      linear-gradient(180deg, #f3fbff 0%, #eaf6ff 45%, #e1f1ff 100%);
    color: #0b3f6b;
}

[data-testid="stAppViewContainer"]::before {
    content: "";
    position: fixed;
    inset: 0;
    background-image: url("data:image/png;base64,__BG_ICON_BASE64__");
    background-repeat: no-repeat;
    background-position: center center;
    background-size: 92vw auto;
    opacity: 0.05;
    pointer-events: none;
    z-index: 0;
}

[data-testid="stHeader"] {
    background: rgba(235, 247, 255, 0.72);
    backdrop-filter: blur(4px);
}

[data-testid="stSidebar"] {
    background: rgba(222, 241, 255, 0.88);
}

h1, h2, h3, h4, p, label, span, div {
    color: #0b3f6b !important;
}

[data-testid="stMetric"] {
    background: rgba(255, 255, 255, 0.82);
    border: 1px solid rgba(33, 122, 181, 0.35);
    border-radius: 14px;
    padding: 10px 12px;
}

[data-testid="stMetricLabel"],
[data-testid="stMetricValue"] {
    color: #0b3f6b !important;
}

[data-testid="stDataFrame"], .stTable {
    border: 1px solid rgba(33, 122, 181, 0.25);
    border-radius: 12px;
    overflow: hidden;
}

[data-testid="stDataFrame"] *,
[data-testid="stDataEditor"] *,
.stTable * {
    color: #0b3f6b !important;
}

[data-testid="stDataEditor"] [role="grid"],
[data-testid="stDataFrame"] [role="grid"] {
    background: rgba(255, 255, 255, 0.9) !important;
}

.stButton > button,
.stDownloadButton > button,
[data-testid="stBaseButton-secondary"] {
    background: linear-gradient(135deg, #1f6aa1, #144e8c) !important;
    color: #ffffff !important;
    border: 1px solid #144e8c !important;
    border-radius: 12px;
    font-weight: 700;
}

.stButton > button,
.stButton > button *,
.stDownloadButton > button,
.stDownloadButton > button *,
[data-testid="stBaseButton-secondary"],
[data-testid="stBaseButton-secondary"] * {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

.stButton > button:hover,
.stDownloadButton > button:hover,
[data-testid="stBaseButton-secondary"]:hover {
    background: linear-gradient(135deg, #1a5c92, #103f73) !important;
    color: #ffffff !important;
    border-color: #103f73 !important;
}

.stButton > button:hover *,
.stDownloadButton > button:hover *,
[data-testid="stBaseButton-secondary"]:hover * {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

div[data-baseweb="select"] > div,
div[data-baseweb="input"] > div,
div[data-baseweb="textarea"] > div,
[data-testid="stNumberInput"] div[data-baseweb="input"] > div {
    background: linear-gradient(135deg, #67afe0, #458ec6) !important;
    border: 1px solid #2b78b7 !important;
}

div[data-baseweb="select"] * ,
div[data-baseweb="input"] input,
div[data-baseweb="textarea"] textarea,
[data-testid="stNumberInput"] input {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

div[data-baseweb="select"] svg {
    fill: #ffffff !important;
}

/* Streamlit top-right toolbar menu */
[data-testid="stToolbar"] [role="menu"],
[data-testid="stToolbar"] [role="menu"] *,
[data-testid="stToolbar"] [role="menuitem"],
[data-testid="stToolbar"] [role="menuitem"] * {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

[data-testid="stToolbar"] [role="menu"] {
    background-color: #0b2f52 !important;
}

/* Select dropdown options must stay readable (light menu + dark text). */
[role="listbox"] {
    background: #f4f9ff !important;
    border: 1px solid #bcd8ef !important;
}

[role="listbox"] * {
    color: #0b3f6b !important;
    -webkit-text-fill-color: #0b3f6b !important;
}

[role="option"] {
    background: #f4f9ff !important;
}

[role="option"][aria-selected="true"] {
    background: #dbeeff !important;
    color: #0b3f6b !important;
}

/* keep content above decorations */
.block-container {
    position: relative;
    z-index: 1;
}

/* Align inputs and buttons to a clean left/right edge */
.stNumberInput, .stSelectbox, .stTextInput, .stTextArea, .stButton, .stDownloadButton {
    width: 100% !important;
}
.stNumberInput > div,
.stSelectbox > div,
.stTextInput > div,
.stTextArea > div,
div[data-baseweb="input"],
div[data-baseweb="select"] {
    width: 100% !important;
    box-sizing: border-box !important;
}
.stButton > button,
.stDownloadButton > button {
    width: 100% !important;
}

/* Additional parameters visual focus */
.extras-card {
    background: rgba(255, 255, 255, 0.72);
    border: 1px solid rgba(33, 122, 181, 0.35);
    border-left: 6px solid #1f6aa1;
    border-radius: 12px;
    padding: 10px 14px;
    margin: 14px 0 10px 0;
    text-align: left !important;
}

.extras-card-title {
    font-weight: 800;
    color: #0b3f6b !important;
    margin: 0 0 2px 0;
    text-align: left !important;
}

.extras-card-note {
    margin: 0;
    color: #2d587e !important;
    text-align: left !important;
}

details[data-testid="stExpander"] summary,
details[data-testid="stExpander"] summary p {
    text-align: left !important;
}

/* One annex field on screen: switches by active tab (1/2 -> balance, 3 -> passport). */
body:not(:has([data-baseweb="tab-list"] button[role="tab"]:nth-child(3)[aria-selected="true"]))
div[data-testid="stTextInput"]:has(input[aria-label="№ приложения по проекту "]) {
    display: none;
}
body:has([data-baseweb="tab-list"] button[role="tab"]:nth-child(3)[aria-selected="true"])
div[data-testid="stTextInput"]:has(input[aria-label="№ приложения по проекту"]) {
    display: none;
}
body:has([data-baseweb="tab-list"] button[role="tab"]:nth-child(3)[aria-selected="true"])
div[data-testid="stTextInput"]:has(input[aria-label="№ приложения по проекту "]) {
    margin-top: -18px !important;
}

/* In Hydraulics tab hide top project meta row (not needed there). */
body:has([data-baseweb="tab-list"] button[role="tab"]:nth-child(2)[aria-selected="true"])
div[data-testid="stSelectbox"]:has([aria-label="Объект"]) {
    display: none !important;
}
body:has([data-baseweb="tab-list"] button[role="tab"]:nth-child(2)[aria-selected="true"])
div[data-testid="stTextInput"]:has(input[aria-label="Наименование"]) {
    display: none !important;
}
body:has([data-baseweb="tab-list"] button[role="tab"]:nth-child(2)[aria-selected="true"])
div[data-testid="stTextInput"]:has(input[aria-label="Адрес"]) {
    display: none !important;
}
body:has([data-baseweb="tab-list"] button[role="tab"]:nth-child(2)[aria-selected="true"])
div[data-testid="stTextInput"]:has(input[aria-label="№ приложения по проекту"]) {
    display: none !important;
}
body:has([data-baseweb="tab-list"] button[role="tab"]:nth-child(2)[aria-selected="true"])
div[data-testid="stTextInput"]:has(input[aria-label="№ приложения по проекту "]) {
    display: none !important;
}
</style>
"""
st.markdown(css_theme.replace("__BG_ICON_BASE64__", BG_ICON_BASE64), unsafe_allow_html=True)

CATALOG_PATH = Path(__file__).resolve().parents[1] / "data" / "consumers_catalog.csv"
PEAK_HOUR_FACTOR = 1.8
DAY_FACTOR = 1.0
RESERVE_FACTOR_WATER = 1.0
LEAKAGE_PERCENT = 0.0
CSV_COLUMNS = [
    "name",
    "unit",
    "count",
    "object_kind",
    "q_u_total_l_day",
    "q_u_hot_l_day",
    "q_hr_total_l_h",
    "q_hr_hot_l_h",
    "q0_total_l_s",
    "q0_total_l_h",
    "q0_spec_l_s",
    "q0_spec_l_h",
    "t_hours",
    "source_doc",
    "source_item",
]


def _infer_object_kind(name: str) -> str:
    n = (name or "").strip().lower()
    production_markers = [
        "завод",
        "цех",
        "фабрик",
        "производ",
        "склад",
        "мастерск",
        "карьер",
        "шахт",
        "котельн",
    ]
    if any(m in n for m in production_markers):
        return "production"
    return "nonproduction"


def _infer_consumer_group(name: str) -> str:
    n = (name or "").strip().lower()
    if any(x in n for x in ["котельн", "теплоснабж", "подпитка"]):
        return "Теплоснабжение"
    if any(x in n for x in ["жилые дома", "квартирного типа", "общежит", "гостиниц", "пансионат", "мотел"]):
        return "Жилье и проживание"
    if any(x in n for x in ["санатор", "дом отдыха"]):
        return "Рекреация и отдых"
    if any(x in n for x in ["бани", "душевые в бытовых помещениях", "душевая кабина", "ванная кабина"]):
        return "Бани и душевые"
    if any(x in n for x in ["предприятия общественного питания", "кафе", "ресторан"]):
        return "Предприятия питания"
    if any(
        x in n
        for x in [
            "дошколь",
            "общеобразователь",
            "образователь",
            "школ",
            "интернат",
            "вуз",
            "преподават",
            "учащ",
        ]
    ):
        return "Образование"
    if any(x in n for x in ["больниц", "поликлиник", "амбулатор", "аптек"]):
        return "Медицина"
    if any(x in n for x in ["плаватель", "спорт", "стадион", "бассейн"]):
        return "Спорт и бассейны"
    if any(x in n for x in ["кинотеатр", "клуб", "театр"]):
        return "Культура и зрелищные"
    if any(x in n for x in ["магазин", "торгов", "парикмахер"]):
        return "Торговля и услуги"
    if any(x in n for x in ["прачечн"]):
        return "Прачечные"
    if any(x in n for x in ["административ"]):
        return "Административные"
    if any(x in n for x in ["полив", "заливка поверхности катка"]):
        return "Полив и благоустройство"
    if any(x in n for x in ["вокзал", "аэропорт", "автовокзал"]):
        return "Транспорт"
    if any(x in n for x in ["цех", "завод", "фабрик", "производ", "промышл"]):
        return "Производство"
    return "Прочее"


def _is_irrigation_consumer(name: str) -> bool:
    n = (name or "").strip().lower()
    return ("полив" in n) or ("заливка поверхности катка" in n)


def _is_no_time_consumer(name: str) -> bool:
    n = (name or "").strip().lower()
    markers = ["полив", "заливка поверхности катка", "подпитка котельной"]
    return any(m in n for m in markers)


def _is_people_unit(unit: str) -> bool:
    u = (unit or "").strip().lower()
    positive = [
        "чел",
        "человек",
        "работающ",
        "учащ",
        "преподав",
        "больной",
        "койка",
        "место",
        "посетител",
        "спортсмен",
        "физкультур",
        "ребенок",
        "артист",
    ]
    negative = ["м2", "м²", "кг", "блюдо", "прибор", "душевая сетка", "%"]
    return any(k in u for k in positive) and not any(k in u for k in negative)


def _can_use_prod_water_source(row: dict, selected_object_kind: str) -> bool:
    if (selected_object_kind or "").strip().lower() != "production":
        return False
    name_l = str(row.get("name", "") or "").strip().lower()
    unit = str(row.get("unit", "") or "")
    if _is_people_unit(unit):
        return False
    disallow_markers = [
        "жил",
        "квартир",
        "сотруд",
        "персонал",
        "офис",
        "гостиниц",
        "общежит",
        "учащ",
        "преподав",
        "больн",
        "пациент",
        "дет",
        "посетител",
    ]
    if any(m in name_l for m in disallow_markers):
        return False
    allow_markers = [
        "цех",
        "производ",
        "технолог",
        "оборуд",
        "станок",
        "мойка",
        "промыв",
        "охлажд",
        "подпитка котельной",
        "котельн",
        "реактор",
        "линия",
        "лаборатор",
        "склад",
    ]
    return any(m in name_l for m in allow_markers)


def _normalize_consumer_row(row: dict) -> dict:
    raw_name = (row.get("name") or "").strip()
    raw_kind = (row.get("object_kind") or "").strip().lower()
    raw_prod_src = row.get("use_prod_water_source", False)
    if isinstance(raw_prod_src, str):
        use_prod_water_source = raw_prod_src.strip().lower() in ("1", "true", "yes", "y", "да")
    else:
        use_prod_water_source = bool(raw_prod_src)
    return {
        "name": raw_name,
        "unit": (row.get("unit") or "").strip(),
        "count": float(row.get("count") or 0.0),
        "use_prod_water_source": use_prod_water_source,
        "np_source_override": (row.get("np_source_override") or "").strip(),
        "np_sewer_override": (row.get("np_sewer_override") or "").strip(),
        "sewer_target_override": (row.get("sewer_target_override") or "").strip(),
        "water_quality_override": (row.get("water_quality_override") or "").strip(),
        "object_kind": raw_kind if raw_kind in ("production", "nonproduction") else _infer_object_kind(raw_name),
        "q_u_total_l_day": float(row.get("q_u_total_l_day") or 0.0),
        "q_u_hot_l_day": float(row.get("q_u_hot_l_day") or 0.0),
        "q_hr_total_l_h": float(row.get("q_hr_total_l_h") or 0.0),
        "q_hr_hot_l_h": float(row.get("q_hr_hot_l_h") or 0.0),
        "q0_total_l_s": float(row.get("q0_total_l_s") or 0.0),
        "q0_total_l_h": float(row.get("q0_total_l_h") or 0.0),
        "q0_spec_l_s": float(row.get("q0_spec_l_s") or 0.0),
        "q0_spec_l_h": float(row.get("q0_spec_l_h") or 0.0),
        "t_hours": float(row.get("t_hours") or 24.0),
        "cold_l_per_unit_day": max(float(row.get("q_u_total_l_day") or 0.0) - float(row.get("q_u_hot_l_day") or 0.0), 0.0),
        "hot_l_per_unit_day": float(row.get("q_u_hot_l_day") or 0.0),
        "source_doc": (row.get("source_doc") or "").strip(),
        "source_item": (row.get("source_item") or "").strip(),
    }


def _load_rows_from_csv_text(text: str) -> list[dict]:
    reader = csv.DictReader(StringIO(text))
    return [_normalize_consumer_row(row) for row in reader]


def _read_catalog_rows() -> list[dict]:
    if not CATALOG_PATH.exists():
        return []
    text = CATALOG_PATH.read_text(encoding="utf-8")
    return _load_rows_from_csv_text(text)


def _build_catalog_index() -> dict[str, dict]:
    return {row.get("name", ""): row for row in st.session_state.get("catalog_rows", [])}


def _autofill_water_rows_from_catalog(rows: list[dict], overwrite: bool) -> list[dict]:
    catalog = _build_catalog_index()
    out: list[dict] = []
    for row in rows:
        normalized = _normalize_consumer_row(row)
        cat = catalog.get(normalized["name"])
        if cat:
            cat_norm = _normalize_consumer_row(cat)
            if overwrite or normalized["cold_l_per_unit_day"] <= 0:
                normalized["cold_l_per_unit_day"] = cat_norm["cold_l_per_unit_day"]
            if overwrite or normalized["hot_l_per_unit_day"] <= 0:
                normalized["hot_l_per_unit_day"] = cat_norm["hot_l_per_unit_day"]
            if overwrite or not normalized["unit"]:
                normalized["unit"] = cat_norm["unit"]
            if overwrite or not normalized["source_doc"]:
                normalized["source_doc"] = cat_norm["source_doc"]
            if overwrite or not normalized["source_item"]:
                normalized["source_item"] = cat_norm["source_item"]
        out.append(normalized)
    return out


def _apply_a2_notes_rules(
    rows: list[dict],
    selected_object_kind: str,
    shift_count: int,
    shift_hours: float,
    apply_shift_rules: bool,
    apartment_rooms_k: int,
    use_apartment_formula: bool,
    use_food_formula: bool,
    food_seats_n: float,
    food_m: float,
    food_t_hours: float,
    food_y: float,
    prod_household_coeff: float,
    laundry_hot_uplift_pct: float,
    use_global_work_hours: bool,
    global_work_hours: float,
) -> list[dict]:
    out: list[dict] = []
    dish_count_hour = 0.0
    dish_count_day = 0.0
    if use_food_formula:
        dish_count_hour = max(0.0, 2.2 * max(food_seats_n, 0.0) * max(food_m, 0.0))
        dish_count_day = dish_count_hour * max(food_t_hours, 0.0) * max(food_y, 0.0)

    for raw in rows:
        row = _normalize_consumer_row(raw)
        name_l = row.get("name", "").strip().lower()
        row["object_kind"] = "production" if selected_object_kind == "production" else "nonproduction"

        if use_global_work_hours and not _is_no_time_consumer(name_l):
            row["t_hours"] = min(24.0, max(float(global_work_hours), 0.0))

        if apply_shift_rules and "в смену" in str(row.get("unit", "")).lower():
            row["count"] = float(row.get("count", 0.0) or 0.0) * max(int(shift_count), 1)
            if not use_global_work_hours:
                row["t_hours"] = min(24.0, max(float(shift_hours), 0.0) * max(int(shift_count), 1))

        if use_apartment_formula and ("жилые дома квартирного типа" in name_l or "жильцы многоквартирного дома" in name_l):
            row["count"] = float(max(apartment_rooms_k, 0) + 1)

        if use_food_formula and "предприятия общественного питания" in name_l:
            row["count"] = dish_count_day
            if float(row.get("t_hours", 0.0) or 0.0) <= 0:
                row["t_hours"] = max(food_t_hours, 1.0)

        if float(prod_household_coeff) != 1.0 and row.get("object_kind") == "production":
            row["count"] = float(row.get("count", 0.0) or 0.0) * max(float(prod_household_coeff), 0.0)

        if laundry_hot_uplift_pct > 0 and "прачечные немеханизированные" in name_l:
            hot_mult = 1.0 + min(max(laundry_hot_uplift_pct, 0.0), 30.0) / 100.0
            row["q_u_hot_l_day"] = float(row.get("q_u_hot_l_day", 0.0) or 0.0) * hot_mult
            row["q_hr_hot_l_h"] = float(row.get("q_hr_hot_l_h", 0.0) or 0.0) * hot_mult
            row["hot_l_per_unit_day"] = float(row.get("hot_l_per_unit_day", 0.0) or 0.0) * hot_mult

        out.append(row)
    return out


def _init_session_state() -> None:
    st.session_state.catalog_rows = _read_catalog_rows()

    if "water_consumers" not in st.session_state:
        st.session_state.water_consumers = []

    # Миграция: удаляем старый автопресет из 2 строк только один раз.
    if not st.session_state.get("water_consumers_default_cleared_v1", False):
        rows = list(st.session_state.get("water_consumers", []))
        if len(rows) == 2:
            names = {str(rows[0].get("name", "")).strip(), str(rows[1].get("name", "")).strip()}
            if names == {"Жильцы многоквартирного дома", "Сотрудники офиса"}:
                st.session_state["water_consumers"] = []
        st.session_state["water_consumers_default_cleared_v1"] = True

def _consumers_to_models(rows: list[dict]) -> list[WaterConsumer]:
    models: list[WaterConsumer] = []
    for row in rows:
        try:
            models.append(
                WaterConsumer(
                    name=str(row.get("name", "")).strip() or "Группа",
                    unit=str(row.get("unit", "")).strip() or "ед",
                    count=float(row.get("count", 0.0)),
                    use_prod_water_source=bool(row.get("use_prod_water_source", False)),
                    cold_l_per_unit_day=float(row.get("cold_l_per_unit_day", 0.0)),
                    hot_l_per_unit_day=float(row.get("hot_l_per_unit_day", 0.0)),
                    q_u_total_l_day=float(row.get("q_u_total_l_day", 0.0)),
                    q_u_hot_l_day=float(row.get("q_u_hot_l_day", 0.0)),
                    q_hr_total_l_h=float(row.get("q_hr_total_l_h", 0.0)),
                    q_hr_hot_l_h=float(row.get("q_hr_hot_l_h", 0.0)),
                    q0_total_l_s=float(row.get("q0_total_l_s", 0.0)),
                    q0_total_l_h=float(row.get("q0_total_l_h", 0.0)),
                    q0_spec_l_s=float(row.get("q0_spec_l_s", 0.0)),
                    q0_spec_l_h=float(row.get("q0_spec_l_h", 0.0)),
                    t_hours=float(row.get("t_hours", 24.0)),
                    source_doc=str(row.get("source_doc", "")).strip(),
                    source_item=str(row.get("source_item", "")).strip(),
                    object_kind=str(row.get("object_kind", "nonproduction")).strip().lower() or "nonproduction",
                    sewer_target_override=str(row.get("sewer_target_override", "")).strip().lower(),
                    water_quality_override=str(row.get("water_quality_override", "")).strip(),
                    np_source_override=str(row.get("np_source_override", "")).strip(),
                    np_sewer_override=str(row.get("np_sewer_override", "")).strip(),
                )
            )
        except Exception:
            continue
    return models


def _file_export_widget(label: str, data: bytes, file_name: str, key: str, mime: str) -> None:
    if IS_NATIVE_APP:
        if st.button(label, use_container_width=True, key=f"{key}_save"):
            try:
                out_dir = Path.home() / "Downloads"
                out_dir.mkdir(parents=True, exist_ok=True)
                target = out_dir / file_name
                if target.exists():
                    stem = target.stem
                    suffix = target.suffix
                    target = out_dir / f"{stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{suffix}"
                target.write_bytes(data)
                st.success(f"Файл сохранен: {target}")
            except Exception as exc:
                st.error(f"Не удалось сохранить файл: {exc}")
        st.caption("Для нативной версии файл сохраняется в папку Downloads.")
    else:
        st.download_button(
            label,
            data=data,
            file_name=file_name,
            mime=mime,
            use_container_width=True,
            key=f"{key}_download",
        )


def _doc_export_widget(label: str, data: bytes, file_name: str, key: str) -> None:
    _file_export_widget(
        label=label,
        data=data,
        file_name=file_name,
        key=key,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _build_hydraulic_sheet_png(
    material_label: str,
    system_label: str,
    scope_label: str,
    speed_mode_label: str,
    q_l_s: float,
    temp_c: float,
    d_in_mm: float,
    length_m: float,
    local_mode_label: str,
    k_local: float,
    xi_sum: float,
    v_limit: float,
    v_m_s: float,
    i_m_per_m: float,
    h_friction_m: float,
    h_local_m: float,
    h_total_m: float,
    re_value: float,
    lambda_f: float,
    nu_m2_s: float,
) -> bytes:
    width, height = 1500, 1060
    img = Image.new("RGB", (width, height), (247, 252, 255))
    draw = ImageDraw.Draw(img)
    font_candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    ]

    def _load_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
        preferred = []
        if bold:
            preferred = [
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
                "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf",
            ]
        else:
            preferred = [
                "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
                "/System/Library/Fonts/Supplemental/Arial.ttf",
                "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
            ]
        for path in preferred + font_candidates:
            if not Path(path).exists():
                continue
            try:
                return ImageFont.truetype(path, size=size)
            except Exception:
                continue
        for path in font_candidates:
            for path in font_candidates:
                if Path(path).exists():
                    try:
                        return ImageFont.truetype(path, size=size)
                    except Exception:
                        continue
        return ImageFont.load_default()

    f_title = _load_font(34, bold=True)
    f_text = _load_font(25, bold=False)
    f_small = _load_font(21, bold=False)
    f_card = _load_font(17, bold=False)
    f_bold = _load_font(27, bold=True)

    draw.rectangle([(20, 20), (width - 20, 180)], fill=(232, 244, 255), outline=(26, 98, 156), width=2)
    scope_l = (scope_label or "").strip().lower()
    if "внутрен" in scope_l:
        sp_ref = "СП 30.13330.2020"
    elif "наруж" in scope_l:
        sp_ref = "СП 31.13330.2021"
    else:
        sp_ref = "СП 30.13330.2020 / СП 31.13330.2021"
    draw.text((38, 34), "Лист гидравлического расчета", fill=(11, 63, 107), font=f_title)
    draw.text(
        (38, 82),
        f"Расчет выполнен по {sp_ref}",
        fill=(11, 63, 107),
        font=f_text,
    )
    draw.text((38, 115), "и таблицам гидравлического расчета Шевелева/Добромыслова", fill=(11, 63, 107), font=f_text)
    draw.text((38, 146), "с учетом параметров производителя труб.", fill=(11, 63, 107), font=f_text)

    draw.rectangle([(20, 200), (width - 20, 540)], fill=(255, 255, 255), outline=(26, 98, 156), width=2)
    left_x = 38
    y = 210
    line_h = 30
    lines = [
        f"Материал труб: {material_label}",
        f"Тип водопровода: {scope_label}",
        f"Система: {system_label}",
        f"Скоростной режим: {speed_mode_label}",
        f"Расход Q: {q_l_s:.3f} л/с",
        f"Температура: {temp_c:.1f} °C",
        f"Расчетный внутренний диаметр: {d_in_mm:.1f} мм",
        f"Длина участка L: {length_m:.2f} м",
        "Учет местных сопротивлений:",
        f"{local_mode_label}",
    ]
    if "коэффициенту k" in local_mode_label.lower():
        lines.append(f"k = {k_local:.3f}")
    if "Σξ" in local_mode_label:
        lines.append(f"Σξ = {xi_sum:.3f}")
    for ln in lines:
        draw.text((left_x, y), ln, fill=(11, 63, 107), font=f_text)
        y += line_h

    # Правая часть блока исходных данных: визуализация потока + карточка участка.
    flow_box = (800, 220, width - 60, 530)
    flow_l, flow_t, flow_r, flow_b = flow_box
    draw.text((flow_l + 12, flow_t + 4), "Ввод / участок сети", fill=(20, 45, 72), font=f_small)
    axis_y = flow_t + 130
    axis_x1 = flow_l + 18
    axis_x2 = flow_r - 120
    flow_color = (48, 132, 195)
    draw.line([(axis_x1, axis_y), (axis_x2, axis_y)], fill=flow_color, width=7)
    arrow_count = 5
    span = (axis_x2 - axis_x1) / (arrow_count + 1)
    for idx in range(arrow_count):
        cx = int(axis_x1 + span * (idx + 1))
        tri = [(cx - 10, axis_y - 58), (cx + 10, axis_y), (cx - 10, axis_y + 58)]
        draw.polygon(tri, fill=flow_color, outline=(27, 101, 162))

    card_x = axis_x2 - 65
    card_y = flow_t + 145
    card_lines = [
        "Карточка участка",
        f"Q = {q_l_s:.3f} л/с",
        f"Dвн = {d_in_mm:.1f} мм",
        f"L = {length_m:.2f} м",
        f"v = {v_m_s:.3f} м/с (предел {v_limit:.2f})",
        f"h_f = {h_friction_m:.3f} м",
        f"h_m = {h_local_m:.3f} м",
        f"H = {h_total_m:.3f} м",
    ]
    cy = card_y
    for i, txt in enumerate(card_lines):
        draw.text((card_x, cy), txt, fill=(20, 45, 72), font=f_card)
        cy += 21 if i == 0 else 20

    draw.rectangle([(20, 550), (width - 20, 920)], fill=(255, 255, 255), outline=(26, 98, 156), width=2)
    draw.text((38, 570), "Результаты расчета", fill=(11, 63, 107), font=f_title)
    ry = 618
    rlines = [
        f"Скорость v: {v_m_s:.3f} м/с (предел: {v_limit:.2f} м/с)",
        f"Гидравлический уклон i: {i_m_per_m:.6f} м/м   |   1000i: {i_m_per_m * 1000.0:.3f} мм/м",
        f"Потери по длине h_f: {h_friction_m:.3f} м",
        f"Местные потери h_m: {h_local_m:.3f} м",
        f"Итого потери напора H: {h_total_m:.3f} м",
        f"Re: {re_value:.0f}",
        f"Коэффициент трения λ: {lambda_f:.6f}",
        f"Кинематическая вязкость ν: {nu_m2_s:.3e} м²/с",
    ]
    for ln in rlines:
        draw.text((45, ry), ln, fill=(11, 63, 107), font=f_text)
        ry += 38

    # Правая часть блока результатов: диаграмма разложения потерь напора.
    chart_l, chart_t, chart_r, chart_b = 840, 640, width - 65, 880
    draw.rectangle([(chart_l, chart_t), (chart_r, chart_b)], outline=(90, 90, 90), width=2, fill=(255, 255, 255))
    chart_title = "Разложение потерь напора на участке"
    draw.text((chart_l + 170, chart_t - 28), chart_title, fill=(45, 45, 45), font=f_small)

    plot_l = chart_l + 28
    plot_r = chart_r - 28
    plot_t = chart_t + 24
    plot_b = chart_b - 48
    draw.line([(plot_l, plot_b), (plot_r, plot_b)], fill=(80, 80, 80), width=2)
    draw.line([(plot_l, plot_t), (plot_l, plot_b)], fill=(80, 80, 80), width=2)
    # Вертикальная подпись оси Y (между рамкой графика и осью).
    y_axis_label = "м вод. ст."
    y_lbl_img = Image.new("RGBA", (220, 40), (255, 255, 255, 0))
    y_lbl_draw = ImageDraw.Draw(y_lbl_img)
    y_lbl_draw.text((0, 0), y_axis_label, fill=(60, 60, 60, 255), font=f_small)
    y_lbl_rot = y_lbl_img.rotate(90, expand=True)
    y_lbl_x = chart_l + 2
    y_lbl_y = (plot_t + plot_b) // 2 - y_lbl_rot.height // 2 - 16
    img.paste(y_lbl_rot, (y_lbl_x, y_lbl_y), y_lbl_rot)

    bar_l = plot_l + 30
    bar_r = plot_r - 30
    bar_h = 44
    bar_b = plot_b - 16
    bar_t = bar_b - bar_h
    hf = max(float(h_friction_m), 0.0)
    hm = max(float(h_local_m), 0.0)
    total = max(float(h_total_m), 0.0)
    frac_hf = (hf / total) if total > 1.0e-12 else 0.0
    split_x = int(bar_l + (bar_r - bar_l) * frac_hf)

    c_hf = (44, 126, 184)
    c_hm = (255, 127, 14)
    draw.rectangle([(bar_l, bar_t), (split_x, bar_b)], fill=c_hf)
    draw.rectangle([(split_x, bar_t), (bar_r, bar_b)], fill=c_hm)
    draw.rectangle([(bar_l, bar_t), (bar_r, bar_b)], outline=(70, 70, 70), width=1)

    pct_hf = int(round((hf / total) * 100.0)) if total > 1.0e-12 else 0
    pct_hm = max(0, 100 - pct_hf)
    if split_x - bar_l > 42:
        draw.text((bar_l + (split_x - bar_l) // 2 - 16, bar_t + 8), f"{pct_hf}%", fill=(255, 255, 255), font=f_small)
    if bar_r - split_x > 42:
        draw.text((split_x + (bar_r - split_x) // 2 - 16, bar_t + 8), f"{pct_hm}%", fill=(20, 20, 20), font=f_small)

    draw.text((bar_l + 250, bar_t - 26), f"Итого H = {total:.3f} м", fill=(60, 60, 60), font=f_small)
    draw.text((bar_l + 272, plot_b + 10), "Потери напора", fill=(60, 60, 60), font=f_small)

    lg_x = plot_r - 215
    lg_y = plot_t + 10
    draw.rectangle([(lg_x, lg_y), (lg_x + 14, lg_y + 14)], fill=c_hf)
    draw.text((lg_x + 24, lg_y - 3), "По длине h_f, м", fill=(55, 55, 55), font=f_small)
    draw.rectangle([(lg_x, lg_y + 28), (lg_x + 14, lg_y + 42)], fill=c_hm)
    draw.text((lg_x + 24, lg_y + 24), "Местные h_m, м", fill=(55, 55, 55), font=f_small)

    ok_text = "Скоростной режим соблюден." if v_m_s <= v_limit else "ВНИМАНИЕ: превышение допустимой скорости."
    ok_color = (21, 128, 61) if v_m_s <= v_limit else (180, 38, 38)
    draw.rectangle([(20, 940), (width - 20, 1020)], fill=(255, 255, 255), outline=(26, 98, 156), width=2)
    draw.text((38, 965), ok_text, fill=ok_color, font=f_bold)

    buf = BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _build_hydraulic_sheet_docx(
    material_label: str,
    system_label: str,
    scope_label: str,
    speed_mode_label: str,
    q_l_s: float,
    temp_c: float,
    d_in_mm: float,
    length_m: float,
    local_mode_label: str,
    k_local: float,
    xi_sum: float,
    v_limit: float,
    v_m_s: float,
    i_m_per_m: float,
    h_friction_m: float,
    h_local_m: float,
    h_total_m: float,
    re_value: float,
    lambda_f: float,
    nu_m2_s: float,
) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    scope_l = (scope_label or "").strip().lower()
    if "внутрен" in scope_l:
        sp_ref = "СП 30.13330.2020"
    elif "наруж" in scope_l:
        sp_ref = "СП 31.13330.2021"
    else:
        sp_ref = "СП 30.13330.2020 / СП 31.13330.2021"

    doc.add_heading("Лист гидравлического расчета", level=1)
    doc.add_paragraph(f"Расчет выполнен по {sp_ref}")
    doc.add_paragraph("и таблицам гидравлического расчета Шевелева/Добромыслова")
    doc.add_paragraph("с учетом параметров производителя труб.")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Параметр"
    hdr[1].text = "Значение"

    rows = [
        ("Материал труб", material_label),
        ("Тип водопровода", scope_label),
        ("Система", system_label),
        ("Скоростной режим", speed_mode_label),
        ("Расход Q", f"{q_l_s:.3f} л/с"),
        ("Температура", f"{temp_c:.1f} °C"),
        ("Расчетный внутренний диаметр", f"{d_in_mm:.1f} мм"),
        ("Длина участка L", f"{length_m:.2f} м"),
        ("Учет местных сопротивлений", local_mode_label),
    ]
    if "коэффициенту k" in local_mode_label.lower():
        rows.append(("k", f"{k_local:.3f}"))
    if "Σξ" in local_mode_label:
        rows.append(("Σξ", f"{xi_sum:.3f}"))
    rows.extend(
        [
            ("Скорость v", f"{v_m_s:.3f} м/с"),
            ("Предел скорости", f"{v_limit:.2f} м/с"),
            ("Гидравлический уклон i", f"{i_m_per_m:.6f} м/м"),
            ("Уклон 1000i", f"{i_m_per_m * 1000.0:.3f} мм/м"),
            ("Потери по длине h_f", f"{h_friction_m:.3f} м"),
            ("Местные потери h_m", f"{h_local_m:.3f} м"),
            ("Итого потери H", f"{h_total_m:.3f} м"),
            ("Re", f"{re_value:.0f}"),
            ("Коэффициент трения λ", f"{lambda_f:.6f}"),
            ("Кинематическая вязкость ν", f"{nu_m2_s:.3e} м²/с"),
        ]
    )

    for name, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(name)
        cells[1].text = str(value)

    doc.add_paragraph("")
    if v_m_s > v_limit:
        doc.add_paragraph("ВНИМАНИЕ: превышение допустимой скорости.")
    else:
        doc.add_paragraph("Скоростной режим соблюден.")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _build_head_meter_calc_docx(
    material_label: str,
    scope_label: str,
    system_label: str,
    h_geo_m: float,
    h_losses_m: float,
    h_free_m: float,
    h_hex_m: float,
    h_inlet_m: float,
    i_m_per_m: float,
    l_inlet_m: float,
    meter_active: dict,
    h_meter_active_m: float,
    h_required_m: float,
    fire_mode: bool,
) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    doc.add_heading("Расчет требуемого напора и потерь в счетчиках", level=1)
    doc.add_paragraph(f"Материал труб: {material_label}")
    doc.add_paragraph(f"Тип водопровода: {scope_label}")
    doc.add_paragraph(f"Система: {system_label}")
    doc.add_paragraph("СП 30.13330.2020: п. 8.27 (требуемый напор), п. 12 (потери в водомерах)")
    doc.add_paragraph(f"Проверка п. 12.16: {'с учетом внутреннего пожаротушения' if fire_mode else 'без внутреннего пожаротушения'}")
    doc.add_paragraph("")

    doc.add_paragraph("Общая формула:")
    doc.add_paragraph("Hтр = Hгеом + Hпотерь + Hсвоб + hсч + Hтепл + Hввод")
    doc.add_paragraph("")

    doc.add_paragraph("")
    doc.add_paragraph("Подбор счетчика по табл. 12.1 и потери в счетчике (п. 12): hсч = S · q²")
    t2 = doc.add_table(rows=1, cols=10)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Система"
    t2.rows[0].cells[1].text = "qсрч (за период), м³/ч"
    t2.rows[0].cells[2].text = "Qэкспл табл.12.1, м³/ч"
    t2.rows[0].cells[3].text = "qрасч, л/с"
    t2.rows[0].cells[4].text = "DN, мм"
    t2.rows[0].cells[5].text = "Тип"
    t2.rows[0].cells[6].text = "S, м/(л/с)^2"
    t2.rows[0].cells[7].text = "hсч, м"
    t2.rows[0].cells[8].text = "Допуск п.12.16, м"
    t2.rows[0].cells[9].text = "Исп. в Hтр"
    rr = t2.add_row().cells
    rr[0].text = system_label
    rr[1].text = f"{float(meter_active.get('q_avg_m3_h', 0.0)):.3f}"
    rr[2].text = f"{float(meter_active.get('q_exp_m3_h', 0.0)):.3f}"
    rr[3].text = f"{float(meter_active.get('q_l_s', 0.0)):.3f}"
    rr[4].text = f"{int(meter_active.get('dn', 0))}"
    rr[5].text = str(meter_active.get("meter_type", "-"))
    rr[6].text = f"{float(meter_active.get('s', 0.0)):.5f}"
    rr[7].text = f"{float(meter_active.get('h_m', 0.0)):.3f}"
    rr[8].text = f"{float(meter_active.get('limit_m', 0.0)):.3f}"
    rr[9].text = "Да"

    doc.add_paragraph("")
    doc.add_paragraph(
        f"Итог: Hтр = {h_geo_m:.3f} + {h_losses_m:.3f} + {h_free_m:.3f} + {h_meter_active_m:.3f} + {h_hex_m:.3f} + {h_inlet_m:.3f} = {h_required_m:.3f} м"
    )
    doc.add_paragraph("")
    doc.add_paragraph("Пояснения по значениям:")
    for line in [
        f"{h_geo_m:.3f} - геометрическая высота Hгеом, м.",
        f"{h_losses_m:.3f} - потери напора по гидравлике участка Hпотерь, м.",
        f"{h_free_m:.3f} - свободный напор Hсвоб, м.",
        f"{h_meter_active_m:.3f} - потери в счетчике hсч, м.",
        f"{h_hex_m:.3f} - потери в теплообменнике Hтепл, м.",
        f"{h_inlet_m:.3f} - потери на вводе Hввод = i·Lввода, м.",
        f"{h_required_m:.3f} - итоговый требуемый напор Hтр, м.",
    ]:
        doc.add_paragraph(line)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


_init_session_state()

st.markdown('<h1>💧 Water<span class="din-script">din</span></h1>', unsafe_allow_html=True)
st.caption("Расчеты выполнены в соответствии с СП 30.13330.2020 и действующей нормативной документацией.")

_instr_left, _instr_right = st.columns([8, 3])
with _instr_right:
    with st.expander("Инструкция по применению", expanded=False):
        st.markdown(
            """
**1. Выберите тип объекта**
- `Непроизводственного назначения` или `Производственного назначения`.
- От выбора зависит форма итогового баланса (Форма 1/Форма 2 по ГОСТ).

**2. Добавьте потребителей**
- В блоке `Добавить потребителя (таблица А.2 СП 30.13330.2020)` выбирайте строки из каталога.
- Нормативные значения подставляются из `СП 30.13330.2020, Приложение А.2`:
  `q_u,m`, `q_hr,u`, `q0`, `q0,hr`, `T`.

**3. Заполните количество**
- `Количество` — число потребителей/единиц (чел, мест, м2, кг и т.д.).
- Для строк “в смену” учитывайте профиль сменности (если включен).

**4. Как считается водопотребление**
- Суточный расход:
  `Qсут = q_u,m * U / 1000` (м3/сут).
- Вероятностно-расчетные пики:
  `q = 5 * q0 * alpha` (л/с),
  `qhr = 0.005 * q0,hr * alpha_hr` (м3/ч).
- `alpha` и `alpha_hr` выбираются по `СП 30.13330.2020, Приложение Б` (B.1/B.2) по `P`, `NP`, `Phr`, `NPhr`.

**5. Спецстроки (полив/подпитка)**
- `Полив` показывается отдельным блоком в расчетной таблице Word.
- Для строк без нормативных `q_hr/q0` пиковые значения не формируются искусственно.
- `Подпитка котельной` рассчитывается автоматически и выводится отдельной строкой.

**6. Итоги**
- Блок `Итог ВС` показывает ХВС/ГВС/Итого и водоотведение.
- Баланс в Word формируется по выбранному типу объекта.

**7. Гидравлика**
- Нажмите вкладку `Гидравлика`.
- Выберите материал трубы и систему (ХВС/ГВС).
- Максимальный секундный расход подхватывается из расчета воды автоматически, но его можно изменить вручную.
- Базовые формулы:
  `v = 4Q/(πd²)`,
  `h_f = i·L`,
  `H = h_f·(1+k)` или `H = h_f + Σξ·v²/(2g)` в зависимости от выбранного режима местных сопротивлений.
- Получаете `v`, `1000i`, потери по длине и местные потери, а также предупреждение при нарушении скоростного режима.

**8. Паспорт ГВС**
- Нажмите вкладку `Паспорт ГВС`.
- Расходы для паспорта ГВС подхватываются автоматически из вкладки баланса; при необходимости включите галочку `Ручной ввод расходов/вероятностей`.
- Базовые формулы:
  `QTh = 1.16 * qh_avg * (th - tc) + Qht`,
  `Qhr,h = 1.16 * qh_max * (th - tc) + Qht`,
  `qcir` и `qh,cir` — по формулам и табличной зависимости `СП 30.13330.2020, Приложение Г`.

**9. Экспорт**
- `Скачать Word-отчет по воде` — итоговый расчет и баланс.
- `Скачать Word-паспорт ГВС` — паспорт ГВС по форме.

**Нормативная база в приложении**
- СП 30.13330.2020 (в т.ч. Приложения А.1, А.2, Б, Г),
- формы баланса в логике ГОСТ Р 21.619-2023.
"""
        )
    with st.expander("О программе", expanded=False):
        st.caption("Разработчик: Диана Белова, проектировщик ВК и НВК.")
        st.caption("Назначение: расчет ВС/ВО, гидравлики и паспорта ГВС по действующей нормативной базе.")

legacy_annex_value = str(st.session_state.get("annex_project", "") or "").strip()
if "annex_balance" not in st.session_state:
    st.session_state["annex_balance"] = legacy_annex_value
if "annex_gvs" not in st.session_state:
    st.session_state["annex_gvs"] = legacy_annex_value

obj_col1, obj_col2, obj_col3, obj_col4 = st.columns([1, 2, 1.75, 0.85])
with obj_col1:
    object_type_label = st.selectbox(
        "Объект",
        options=["Непроизводственного назначения", "Производственного назначения"],
        index=0,
    )
with obj_col2:
    object_name_input = st.text_input("Наименование", value=st.session_state.get("object_name_input", ""))
    st.session_state.object_name_input = object_name_input
with obj_col3:
    object_address_input = st.text_input("Адрес", value=st.session_state.get("object_address_input", ""))
    st.session_state.object_address_input = object_address_input
with obj_col4:
    st.text_input(
        "№ приложения по проекту",
        key="annex_balance",
    )
    st.text_input(
        "№ приложения по проекту\u00A0",
        key="annex_gvs",
    )
selected_object_kind = "production" if object_type_label == "Производственного назначения" else "nonproduction"
object_name = (object_name_input or "").strip()
object_address = (object_address_input or "").strip()
project_name = "Расчет водоснабжения и водоотведения"
organization = ""
author = ""
stage = ""
revision = ""

tab_water, tab_hyd, tab_gvs = st.tabs(["1) Баланс ВС и ВО", "2) Гидравлика", "3) Паспорт ГВС"])

with tab_water:
    st.markdown("*Примечание: заполняется по выбранной форме Приложения А ГОСТ Р 21.619-2023.*")
    annex_balance = str(st.session_state.get("annex_balance", "") or "").strip()

    use_apartment_formula = False
    apartment_rooms_k = 1
    use_food_formula = False
    food_seats_n = 0.0
    food_m = 2.0
    food_t_hours = 0.0
    food_y = 0.45
    prod_household_coeff = 1.0
    laundry_hot_uplift_pct = 0.0
    shift_count = 1
    shift_hours = 0.0
    apply_shift_rules = False
    use_global_work_hours = False
    global_work_hours = 0.0
    np_water_source = "Горводопровод"
    np_storm_m3_day = 0.0
    pr_water_source = "Из хоз.-питьевого водопровода"
    pr_concentration_mg_l = ""

    add_col1, add_col2, add_col3 = st.columns(3)
    with add_col1:
        options = ["-"]
        for row in st.session_state.catalog_rows:
            nm = row.get("name", "")
            if nm:
                options.append(f"{_infer_consumer_group(nm)} | {nm}")
        category_to_add = st.selectbox(
            "Добавить потребителя (таблица А.2 СП 30.13330.2020)",
            options=sorted(options),
            label_visibility="collapsed",
        )
    with add_col2:
        if st.button("Добавить", use_container_width=True) and category_to_add != "-":
            selected_name = category_to_add.split(" | ", 1)[1] if " | " in category_to_add else category_to_add
            cat = _build_catalog_index().get(selected_name)
            if cat:
                new_row = _normalize_consumer_row(cat)
                new_row["count"] = 0.0
                st.session_state.water_consumers = list(st.session_state.water_consumers) + [new_row]
    with add_col3:
        if st.button("Очистить таблицу", use_container_width=True):
            st.session_state.water_consumers = []

    st.subheader("Таблица потребителей")
    display_rows = []
    for r in st.session_state.water_consumers:
        rr = dict(r)
        rr["__delete"] = False
        display_rows.append(rr)
    edited_consumers = st.data_editor(
        display_rows,
        num_rows="dynamic",
        use_container_width=True,
        key="water_consumers_editor_v2",
        column_order=[
            "__delete",
            "name",
            "unit",
            "count",
            "sewer_target_override",
            "water_quality_override",
            "q_u_total_l_day",
            "q_u_hot_l_day",
            "q_hr_total_l_h",
            "q_hr_hot_l_h",
            "q0_total_l_s",
            "q0_total_l_h",
            "q0_spec_l_s",
            "q0_spec_l_h",
            "t_hours",
            "source_doc",
            "source_item",
        ],
        column_config={
            "__delete": st.column_config.CheckboxColumn("Выбор"),
            "name": st.column_config.TextColumn("Группа потребителей"),
            "unit": st.column_config.TextColumn("Ед."),
            "count": st.column_config.NumberColumn("Количество", min_value=0.0, step=1.0, format="%.2f"),
            "sewer_target_override": st.column_config.SelectboxColumn(
                "ВО (override)",
                options=["", "domestic", "production"],
                help="Пусто = авто; domestic = бытовая; production = производственная",
            ),
            "water_quality_override": st.column_config.SelectboxColumn(
                "Качество воды (override)",
                options=[
                    "",
                    "Питьевая (СанПиН 2.1.3684-21, 1.2.3685-21)",
                    "Техническая",
                    "Оборотная",
                    "По ТЗ ТХ",
                    "Привозная",
                    "Дождевая",
                    "Производственная",
                ],
                help="Пусто = авто по типу потребителя; можно задать вручную",
            ),
            "q_u_total_l_day": st.column_config.NumberColumn("q_u,m общий, л/ед·сут", min_value=0.0, step=0.1, format="%.2f"),
            "q_u_hot_l_day": st.column_config.NumberColumn("q_u,m горячей, л/ед·сут", min_value=0.0, step=0.1, format="%.2f"),
            "q_hr_total_l_h": st.column_config.NumberColumn("q_hr,u общий, л/ед·ч", min_value=0.0, step=0.1, format="%.2f"),
            "q_hr_hot_l_h": st.column_config.NumberColumn("q_hr,u горячей, л/ед·ч", min_value=0.0, step=0.1, format="%.2f"),
            "q0_total_l_s": st.column_config.NumberColumn("q0 tot, л/с", min_value=0.0, step=0.01, format="%.2f"),
            "q0_total_l_h": st.column_config.NumberColumn("q0hr tot, л/ч", min_value=0.0, step=1.0, format="%.0f"),
            "q0_spec_l_s": st.column_config.NumberColumn("q0 (ХВС/ГВС), л/с", min_value=0.0, step=0.01, format="%.2f"),
            "q0_spec_l_h": st.column_config.NumberColumn("q0hr (ХВС/ГВС), л/ч", min_value=0.0, step=1.0, format="%.0f"),
            "t_hours": st.column_config.NumberColumn("T, ч", min_value=0.0, step=1.0, format="%.0f"),
            "source_doc": st.column_config.TextColumn("Нормативный документ"),
            "source_item": st.column_config.TextColumn("Пункт/таблица"),
        },
    )
    if selected_object_kind == "nonproduction":
        st.caption("*Примечание: отметка в колонке «Выбор» используется не только для удаления строк, но и для применения дополнительных параметров к отмеченным потребителям.*")
    del_col1, del_col2 = st.columns([2, 5])
    with del_col1:
        delete_marked = st.button("Удалить отмеченные строки", use_container_width=True)
    apply_np_overrides = False
    np_override_source_selected = "Авто"
    np_override_sewer_selected = "Авто"
    marked_rows_present = any(bool(src.get("__delete", False)) for src in edited_consumers)
    if selected_object_kind == "nonproduction" and marked_rows_present:
        st.markdown("**Параметры для отмеченных строк (форма 2)**")
        nh1, nh2, nh3 = st.columns([2, 2, 1.5])
        with nh1:
            st.caption("Источник водоснабжения (для отмеченных)")
        with nh2:
            st.caption("Водоотведение (для отмеченных)")
        with nh3:
            st.caption(" ")
        npo1, npo2, npo3 = st.columns([2, 2, 1.5])
        with npo1:
            np_override_source_selected = st.selectbox(
                "Источник водоснабжения (для отмеченных)",
                options=["Авто", "Горводопровод", "Скважины", "Техвода", "Оборотные системы"],
                index=0,
                label_visibility="collapsed",
                key="np_override_source_selected",
            )
        with npo2:
            np_override_sewer_selected = st.selectbox(
                "Водоотведение (для отмеченных)",
                options=["Авто", "Хоз.-быт.", "Норм.-чистые", "Загр. мех./мин.", "Загр. хим./орг."],
                index=0,
                label_visibility="collapsed",
                key="np_override_sewer_selected",
            )
        with npo3:
            apply_np_overrides = st.button("Применить к отмеченным", use_container_width=True)
    normalized_rows = _autofill_water_rows_from_catalog(edited_consumers, overwrite=False)
    if delete_marked:
        st.session_state.water_consumers = []
        for row, src in zip(normalized_rows, edited_consumers):
            if bool(src.get("__delete", False)):
                continue
            clean_row = dict(row)
            clean_row.pop("__delete", None)
            st.session_state.water_consumers.append(clean_row)
        st.rerun()
    elif apply_np_overrides and selected_object_kind == "nonproduction":
        updated_rows = []
        for row, src in zip(normalized_rows, edited_consumers):
            clean_row = dict(row)
            clean_row.pop("__delete", None)
            if bool(src.get("__delete", False)):
                clean_row["np_source_override"] = "" if np_override_source_selected == "Авто" else np_override_source_selected
                clean_row["np_sewer_override"] = "" if np_override_sewer_selected == "Авто" else np_override_sewer_selected
            updated_rows.append(clean_row)
        st.session_state.water_consumers = updated_rows
        st.success("Параметры для отмеченных строк обновлены.")
        st.rerun()
    else:
        cleaned_rows = []
        for row in normalized_rows:
            clean_row = dict(row)
            clean_row.pop("__delete", None)
            cleaned_rows.append(clean_row)
        st.session_state.water_consumers = cleaned_rows

    # Подача из пром. водопровода доступна только для релевантных производственных потребителей.
    enforce_rows = [dict(r) for r in st.session_state.water_consumers]
    for row in enforce_rows:
        if not _can_use_prod_water_source(row, selected_object_kind):
            row["use_prod_water_source"] = False
    st.session_state.water_consumers = enforce_rows

    if selected_object_kind == "production":
        eligible_prod_rows = []
        for i, row in enumerate(st.session_state.water_consumers):
            if _can_use_prod_water_source(row, selected_object_kind):
                eligible_prod_rows.append(
                    {
                        "row_idx": int(i),
                        "consumer": str(row.get("name", "")).strip() or "Группа",
                        "unit": str(row.get("unit", "")).strip(),
                        "use_prod_water_source": bool(row.get("use_prod_water_source", False)),
                    }
                )
        if eligible_prod_rows:
            st.markdown("**Пром. водопровод (только допустимые потребители)**")
            edited_prod_rows = st.data_editor(
                eligible_prod_rows,
                use_container_width=True,
                hide_index=True,
                key="water_prod_source_editor",
                column_order=["consumer", "unit", "use_prod_water_source"],
                disabled=["consumer", "unit"],
                column_config={
                    "consumer": st.column_config.TextColumn("Группа потребителей"),
                    "unit": st.column_config.TextColumn("Ед."),
                    "use_prod_water_source": st.column_config.CheckboxColumn(
                        "Пром. водопровод",
                        help="Отметьте только для строк, где технически допускается подача из производственного водопровода.",
                    ),
                },
            )
            synced_rows = [dict(r) for r in st.session_state.water_consumers]
            for item in edited_prod_rows:
                idx = int(item.get("row_idx", -1))
                if 0 <= idx < len(synced_rows):
                    synced_rows[idx]["use_prod_water_source"] = bool(item.get("use_prod_water_source", False))
            st.session_state.water_consumers = synced_rows

    has_shift_units = any(
        "в смену" in str(r.get("unit", "")).lower()
        for r in st.session_state.water_consumers
        if float(r.get("count", 0.0) or 0.0) > 0 or str(r.get("unit", "")).strip()
    )

    st.markdown(
        """
        <div class="extras-card">
          <p class="extras-card-title">Дополнительные параметры</p>
          <p class="extras-card-note">Настройте дополнительные параметры и значения.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.expander("Открыть/скрыть дополнительные параметры", expanded=False):
        st.markdown("**Режим работы и сменность**")
        if has_shift_units:
            sh1, sh2, sh3 = st.columns(3)
            with sh1:
                shift_preset = st.selectbox(
                    "Профиль сменности",
                    options=["1 смена x 8 ч", "2 смены x 4 ч", "1 смена x 12 ч", "Пользовательский"],
                    index=3,
                )
            preset_map = {
                "1 смена x 8 ч": (1, 8.0),
                "2 смены x 4 ч": (2, 4.0),
                "1 смена x 12 ч": (1, 12.0),
            }
            p_shifts, p_hours = preset_map.get(shift_preset, (1, 0.0))
            with sh2:
                shift_count = st.number_input("Количество смен", min_value=1, max_value=4, value=int(p_shifts), step=1, disabled=(shift_preset != "Пользовательский"))
            with sh3:
                shift_hours = st.number_input("Длительность смены, ч", min_value=0.0, max_value=24.0, value=float(p_hours), step=0.5, disabled=(shift_preset != "Пользовательский"))
            apply_shift_rules = st.checkbox("Автоприменять к строкам с единицей 'в смену'", value=False)
            total_shift_hours = min(24.0, float(shift_count) * float(shift_hours))
        else:
            st.info("В выбранных потребителях нет единиц измерения 'в смену'.")
            total_shift_hours = 0.0
            shift_count = 1
            shift_hours = 0.0
            apply_shift_rules = False
        g1, g2 = st.columns(2)
        with g1:
            use_global_work_hours = st.checkbox("Единое время работы объекта", value=False)
        with g2:
            global_work_hours = st.number_input(
                "Время работы объекта, ч/сут",
                min_value=0.0,
                max_value=24.0,
                value=float(total_shift_hours),
                step=0.5,
                disabled=not bool(use_global_work_hours),
            )
        if selected_object_kind == "nonproduction":
            v1, v2 = st.columns([1, 2])
            with v1:
                np_storm_m3_day = st.number_input("Водосток, м³/сут", min_value=0.0, value=0.0, step=0.1)
        else:
            v1, v2 = st.columns([1, 2])
            with v1:
                pr_inlet_pressure_mpa = st.number_input(
                    "Расчетное давление на вводе, МПа",
                    min_value=0.0,
                    value=float(st.session_state.get("pr_inlet_pressure_mpa", 0.0) or 0.0),
                    step=0.01,
                    key="pr_inlet_pressure_mpa",
                )

        st.markdown("**Примечания к таблице А.2 СП 30.13330.2020**")
        r1, r2, r3 = st.columns(3)
        with r1:
            use_apartment_formula = st.checkbox("Жилые дома: N = K + 1", value=False)
            apartment_rooms_k = st.number_input("K (число жилых комнат)", min_value=0, value=1, step=1)
        with r2:
            use_food_formula = st.checkbox("Общепит: Uч=2.2*n*m; Uсут=Uч*T*y", value=False)
            food_seats_n = st.number_input("n (посадочных мест)", min_value=0.0, value=0.0, step=1.0)
            food_m = st.number_input("m (посадок)", min_value=0.0, value=2.0, step=0.1)
        with r3:
            food_t_hours = st.number_input("T (часы работы)", min_value=0.0, value=0.0, step=0.5)
            food_y = st.number_input("y (коэф. неравномерности)", min_value=0.0, value=0.45, step=0.01)
            prod_household_coeff = st.number_input("Коэф. 0.6 для производств", min_value=0.0, value=1.0, step=0.1)
            laundry_hot_uplift_pct = st.number_input("Прачечные немех.: +ГВС, % (до 30)", min_value=0.0, max_value=30.0, value=0.0, step=1.0)
        st.caption("Параметры и форма баланса соответствуют структуре ГОСТ Р 21.619-2023 (приложение А).")

    irrigation_rows = [r for r in st.session_state.water_consumers if _is_irrigation_consumer(str(r.get("name", "")))]
    if irrigation_rows:
        with st.expander("Полив: качество воды", expanded=False):
            irrigation_quality = st.selectbox(
                "Качество воды для строк полива",
                options=["Питьевая", "Привозная", "Дождевая", "Производственная"],
                index=0,
            )
            if st.button("Применить качество к поливу", use_container_width=False):
                updated = []
                for row in st.session_state.water_consumers:
                    if _is_irrigation_consumer(str(row.get("name", ""))):
                        row = dict(row)
                        row["water_quality_override"] = irrigation_quality
                    updated.append(row)
                st.session_state.water_consumers = updated
                st.success("Качество воды для полива обновлено.")

    calc_rows = _apply_a2_notes_rules(
        rows=st.session_state.water_consumers,
        selected_object_kind=selected_object_kind,
        shift_count=int(shift_count),
        shift_hours=float(shift_hours),
        apply_shift_rules=bool(apply_shift_rules),
        apartment_rooms_k=int(apartment_rooms_k),
        use_apartment_formula=bool(use_apartment_formula),
        use_food_formula=bool(use_food_formula),
        food_seats_n=float(food_seats_n),
        food_m=float(food_m),
        food_t_hours=float(food_t_hours),
        food_y=float(food_y),
        prod_household_coeff=float(prod_household_coeff),
        laundry_hot_uplift_pct=float(laundry_hot_uplift_pct),
        use_global_work_hours=bool(use_global_work_hours),
        global_work_hours=float(global_work_hours),
    )
    water_models = _consumers_to_models(calc_rows)
    water_res = calc_water_by_consumers_advanced(
        consumers=water_models,
        peak_hour_factor=PEAK_HOUR_FACTOR,
        day_factor=DAY_FACTOR,
        reserve_factor=RESERVE_FACTOR_WATER,
        leakage_percent=LEAKAGE_PERCENT,
    )
    detected_form = "Форма 1 (производственный объект)" if selected_object_kind == "production" else "Форма 2 (непроизводственный объект)"
    st.caption(f"Автоопределение типа объекта для итогового Word: {detected_form}")

    st.subheader("Итог ВС")
    m1, m2, m3, m4, m5 = st.columns(5)
    m1.metric("ХВС, м³/сут", f'{water_res["cold_m3_day"]:.3f}')
    m2.metric("ГВС, м³/сут", f'{water_res["hot_m3_day"]:.3f}')
    m3.metric("Итого, м³/сут", f'{water_res["total_m3_day"]:.3f}')
    m4.metric("Макс. час, м³/ч", f'{water_res["max_m3_hour"]:.3f}')
    m5.metric("Макс. расход, л/с", f'{water_res["max_l_sec"]:.3f}')

    if water_res["rows"]:
        water_df = pd.DataFrame(water_res["rows"])
        st.dataframe(water_df, use_container_width=True, hide_index=True)

    st.subheader("Итоговый баланс водоснабжения и водоотведения")
    st.caption("Форма баланса выполнена в логике таблиц СПДС (ГОСТ Р 21.101-2020) и структуры баланса ГОСТ Р 21.619-2023 (приложение А).")
    balance_df = pd.DataFrame(water_res.get("balance_rows", []))
    if not balance_df.empty:
        st.dataframe(balance_df, use_container_width=True, hide_index=True)

    gvs_res_for_report = calc_gvs_passport(
        qh_avg_m3_h=float(water_res["hot_avg_m3_hour"]),
        qh_max_m3_h=float(water_res["hot_max_m3_hour"]),
        t_hot_c=float(st.session_state.get("t_hot_c", 60.0)),
        t_cold_c=float(st.session_state.get("t_cold_c", 10.0)),
        qht_kW=float(st.session_state.get("qht_kw", 2.88)),
        delta_t_supply_c=float(st.session_state.get("delta_t_supply_c", 10.0)),
    )
    checks = build_data_checks(water_res["rows"], [], require_heat_elements=False)
    if float(water_res["max_l_sec"]) <= 0.0:
        checks.append("Максимальный расчетный расход воды равен 0 л/с.")

    project_meta = {
        "organization": organization,
        "author": author,
        "stage": stage,
        "revision": revision,
    }
    water_inputs_for_doc = {
        "peak_hour_factor": f"{PEAK_HOUR_FACTOR:.2f}",
        "day_factor": f"{DAY_FACTOR:.2f}",
        "reserve_factor": f"{RESERVE_FACTOR_WATER:.2f}",
        "leakage_percent": f"{LEAKAGE_PERCENT:.2f}",
        "adjustment_factor": f'{float(water_res["adjustment_factor"]):.3f}',
        "selected_object_kind": selected_object_kind,
        "np_water_source": np_water_source if selected_object_kind == "nonproduction" else "",
        "np_storm_m3_day": f"{float(np_storm_m3_day):.3f}" if selected_object_kind == "nonproduction" else "0",
        "pr_water_source": pr_water_source if selected_object_kind == "production" else "",
        "pr_concentration_mg_l": pr_concentration_mg_l if selected_object_kind == "production" else "",
        "pr_inlet_pressure_mpa": f'{float(st.session_state.get("pr_inlet_pressure_mpa", 0.0) or 0.0):.3f}' if selected_object_kind == "production" else "",
        "hyd_required_head_m_hvs": f'{float(st.session_state.get("hyd_required_head_m_hvs", 0.0) or 0.0):.3f}',
        "hyd_required_pressure_mpa": f'{float(st.session_state.get("hyd_required_pressure_mpa", 0.0) or 0.0):.3f}',
        "passport_h_top_m": f'{float(st.session_state.get("passport_h_top", 0.0) or 0.0):.3f}',
        "passport_free_head_m": f'{float(st.session_state.get("passport_free_head_m", 20.0) or 20.0):.3f}',
        "passport_losses_system_m": f'{float(st.session_state.get("passport_losses_system_m", 0.0) or 0.0):.3f}',
        "passport_circ_losses_m": f'{float(st.session_state.get("passport_circ_losses_m", 0.0) or 0.0):.3f}',
        "passport_has_meter": "1" if bool(st.session_state.get("passport_has_meter", True)) else "0",
        "passport_meter_loss_m": f'{float(st.session_state.get("passport_meter_loss_m", 0.0) or 0.0):.3f}',
    }
    report_doc = build_report_docx(
        project_name=project_name,
        object_name=object_name,
        object_address=object_address,
        annex_label=(annex_balance or "").strip(),
        project_meta=project_meta,
        water_inputs=water_inputs_for_doc,
        water_results=water_res,
        water_consumers=water_res["rows"],
        gvs_results=gvs_res_for_report,
        checks=checks,
    )
    _doc_export_widget(
        label="⬇️ Скачать Word-отчет по воде",
        data=report_doc,
        file_name=f"otchet_water_balance_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        key="water_report",
    )

with tab_hyd:
    st.subheader("Гидравлика")
    st.caption(
        "Расчет потерь напора, скоростей, уклонов и подбора диаметра. "
        "Формулы и ограничения применяются для ХВС/ГВС по СП 30.13330.2020, СП 31.13330.2021 и профильным таблицам гидравлики."
    )

    def _parse_sdr(label: str) -> float | None:
        m = re.search(r"SDR\s*([0-9]+(?:\.[0-9]+)?)", str(label))
        return float(m.group(1)) if m else None

    def _reset_hyd_form(material_code: str) -> None:
        keys_to_drop = [
            f"hyd_system_{material_code}",
            f"hyd_temp_{material_code}",
            f"hyd_speed_mode_{material_code}",
            f"hyd_speed_mode_outer_{material_code}",
            f"hyd_q_{material_code}",
            f"hyd_is_new_{material_code}",
            f"hyd_manual_d_chk_{material_code}",
            f"hyd_din_manual_{material_code}",
            f"hyd_din_manual_pending_{material_code}",
            f"hyd_manual_hint_{material_code}",
            f"hyd_dout_{material_code}",
            f"hyd_s_{material_code}",
            f"hyd_din_{material_code}",
            f"hyd_len_{material_code}",
            f"hyd_pe_grade_{material_code}",
            f"hyd_sdr_{material_code}",
            f"hyd_local_mode_{material_code}",
            f"hyd_k_preset_{material_code}",
            f"hyd_k_manual_{material_code}",
            f"hyd_xi_{material_code}",
        ]
        for k in keys_to_drop:
            st.session_state.pop(k, None)

    def _reset_hyd_head_meter_form(material_code: str) -> None:
        keys_to_drop = [
            f"hyd_head_geo_{material_code}",
            f"hyd_head_free_{material_code}",
            f"hyd_head_extra_{material_code}",
            f"hyd_inlet_len_{material_code}",
            f"hyd_meter_fire_{material_code}",
            f"hyd_meter_q_fire_{material_code}",
            f"hyd_meter_qh_hvs_{material_code}",
            f"hyd_meter_q_hvs_{material_code}",
            f"hyd_meter_qh_gvs_{material_code}",
            f"hyd_meter_q_gvs_{material_code}",
        ]
        for k in keys_to_drop:
            st.session_state.pop(k, None)

    def _sync_autofill_number(key: str, auto_value: float) -> None:
        prev_key = f"{key}__auto_prev"
        auto_val = float(auto_value)
        if key not in st.session_state:
            st.session_state[key] = auto_val
            st.session_state[prev_key] = auto_val
            return
        cur_val = float(st.session_state.get(key, 0.0) or 0.0)
        prev_auto = st.session_state.get(prev_key, None)
        if prev_auto is None or abs(cur_val - float(prev_auto)) < 1.0e-12:
            st.session_state[key] = auto_val
        st.session_state[prev_key] = auto_val

    top_c1, top_c2 = st.columns(2)
    with top_c1:
        mat_code = st.selectbox(
            "Материал труб",
            options=list(MATERIALS.keys()),
            format_func=lambda x: MATERIALS.get(x, {}).get("label", x),
            key="hyd_material",
        )
    with top_c2:
        hyd_scope = st.selectbox(
            "Тип водопровода",
            options=["Внутренний", "Наружный"],
            key="hyd_scope",
        )

    mat_doc = MATERIALS.get(mat_code, {}).get("doc", "")
    if mat_doc:
        st.caption(f"Нормативная база: {mat_doc}")

    with st.expander(f"Настройка: {MATERIALS.get(mat_code, {}).get('label', mat_code)}", expanded=True):
        c_top1, c_top2, c_top3, c_top4 = st.columns([1, 1, 1, 1])
        q_key = f"hyd_q_{mat_code}"
        speed_mode_map_inner = {
            "Жилые и общественные здания": 1.5,
            "Административно-бытовые здания предприятий": 2.0,
            "Производственные здания и сооружения": 3.0,
        }
        speed_mode_map_outer = {
            "Наружная сеть населенного пункта": 2.0,
            "Напорные линии насосных станций": 2.5,
            "Наружная противопожарная сеть": 5.0,
        }
        speed_default_idx = 2 if selected_object_kind == "production" else 0
        with c_top1:
            if hyd_scope == "Наружный":
                hyd_system = st.selectbox("Система", options=["ХВС"], key=f"hyd_system_{mat_code}")
                st.caption("Для наружного водопровода в этом разделе рассчитывается только ХВС.")
            else:
                hyd_system = st.selectbox("Система", options=["ХВС", "ГВС"], key=f"hyd_system_{mat_code}")
        with c_top2:
            temp_default = 10.0 if hyd_system == "ХВС" else 60.0
            hyd_temp_c = st.number_input(
                "Температура, °C",
                min_value=1.0,
                max_value=90.0,
                value=float(temp_default),
                step=1.0,
                key=f"hyd_temp_{mat_code}",
            )
        with c_top3:
            if hyd_scope == "Наружный":
                hyd_speed_label = st.selectbox(
                    "Скоростной режим (СП 31.13330.2021)",
                    options=list(speed_mode_map_outer.keys()),
                    index=0,
                    key=f"hyd_speed_mode_outer_{mat_code}",
                )
                v_limit = float(speed_mode_map_outer[hyd_speed_label])
            else:
                hyd_speed_label = st.selectbox(
                    "Скоростной режим (СП 30.13330.2020, п. 8.26)",
                    options=list(speed_mode_map_inner.keys()),
                    index=speed_default_idx,
                    key=f"hyd_speed_mode_{mat_code}",
                )
                v_limit = float(speed_mode_map_inner[hyd_speed_label])
        auto_q_default = 0.0
        if "water_res" in locals():
            if hyd_system == "ГВС":
                auto_q_default = float(water_res.get("hot_max_l_sec", 0.0) or 0.0)
            else:
                auto_q_default = float(water_res.get("cold_max_l_sec", 0.0) or float(water_res.get("max_l_sec", 0.0) or 0.0))
        _sync_autofill_number(q_key, auto_q_default)
        q_default = float(st.session_state.get(q_key, auto_q_default) or auto_q_default)
        with c_top4:
            hyd_q_l_s = st.number_input(
                "Q, л/с",
                min_value=0.0,
                value=q_default,
                step=0.01,
                key=q_key,
                help="Ручной ввод расхода.",
            )
        geom_c1, geom_c2, geom_c3, geom_c4 = st.columns([1, 1, 1, 1])
        manual_key = f"hyd_manual_d_chk_{mat_code}"
        manual_hint_key = f"hyd_manual_hint_{mat_code}"
        manual_value_key = f"hyd_din_manual_{mat_code}"
        manual_pending_key = f"hyd_din_manual_pending_{mat_code}"
        pending_manual_val = st.session_state.pop(manual_pending_key, None)
        base_manual_default = 0.0 if manual_value_key not in st.session_state else float(st.session_state.get(manual_value_key, 0.0) or 0.0)
        manual_default_val = float(pending_manual_val) if pending_manual_val is not None else base_manual_default
        with geom_c1:
            hyd_is_new = st.checkbox("Трубы новые", value=True, key=f"hyd_is_new_{mat_code}")
            use_manual_din = st.checkbox("Ручной ввод dвн", value=False, key=manual_key)

        d_out_mm = 0.0
        s_mm = 0.0
        d_in_mm = 0.0
        len_key = f"hyd_len_{mat_code}"
        len_default = 0.0 if len_key not in st.session_state else float(st.session_state[len_key] or 0.0)
        hyd_len_m = float(len_default)

        if mat_code in ("steel_vgp", "steel_welded", "plastic", "metal_plastic", "polyplastic", "copper"):
            if mat_code in ("steel_vgp", "steel_welded"):
                dims_map = STEEL_DIMENSIONS
            elif mat_code == "plastic":
                dims_map = PLASTIC_DIMENSIONS
            elif mat_code == "metal_plastic":
                dims_map = METAL_PLASTIC_DIMENSIONS
            elif mat_code == "copper":
                dims_map = COPPER_DIMENSIONS
            else:
                dims_map = POLYPLASTIC_DIMENSIONS
            with geom_c2:
                if use_manual_din:
                    st.number_input("Наружный диаметр, мм", min_value=0.0, value=0.0, step=0.1, disabled=True, key=f"hyd_dout_disabled_{mat_code}")
                else:
                    d_out_sel = st.selectbox("Наружный диаметр, мм", options=[0.0] + sorted(dims_map.keys()), key=f"hyd_dout_{mat_code}")
                    d_out_mm = float(d_out_sel)
            with geom_c3:
                if use_manual_din:
                    st.number_input("Толщина стенки, мм", min_value=0.0, value=0.0, step=0.1, disabled=True, key=f"hyd_s_disabled_{mat_code}")
                else:
                    selected_dout = int(d_out_mm)
                    if mat_code in ("plastic", "metal_plastic", "polyplastic"):
                        if mat_code == "plastic":
                            sdr_sel = st.selectbox("Серия SDR/SN", options=PLASTIC_SDR_SERIES, key=f"hyd_sdr_{mat_code}")
                        else:
                            sdr_sel = st.selectbox("Серия SDR/SN", options=MLPEX_SDR_SERIES, key=f"hyd_sdr_{mat_code}")
                        sdr_value = _parse_sdr(sdr_sel)
                        if selected_dout > 0 and sdr_value and sdr_value > 0:
                            s_mm = round(float(selected_dout) / float(sdr_value), 2)
                        else:
                            s_mm = 0.0
                        st.caption(f"Толщина стенки s = {s_mm:.2f} мм (авто)")
                    elif mat_code == "copper":
                        selected_dout_float = float(d_out_mm)
                        s_options = [0.0] if selected_dout_float <= 0.0 else dims_map.get(selected_dout_float, [0.0])
                        s_sel = st.selectbox("Толщина стенки, мм", options=s_options, key=f"hyd_s_{mat_code}")
                        s_mm = float(s_sel)
                    else:
                        s_sel = st.selectbox("Толщина стенки, мм", options=[0.0] if selected_dout == 0 else dims_map[selected_dout], key=f"hyd_s_{mat_code}")
                        s_mm = float(s_sel)
                    d_in_mm = max(d_out_mm - 2.0 * s_mm, 0.0)
        elif mat_code == "cast_iron":
            with geom_c2:
                if use_manual_din:
                    st.number_input(
                        "DN (номинальный диаметр), мм",
                        min_value=0.0,
                        value=0.0,
                        step=1.0,
                        disabled=True,
                        key=f"hyd_dn_disabled_{mat_code}",
                    )
                    dn_sel = 0
                    cast_class = st.selectbox(
                        "Класс давления",
                        options=list(CAST_IRON_BY_CLASS.keys()),
                        index=0,
                        disabled=True,
                        key=f"hyd_cast_class_{mat_code}",
                    )
                else:
                    cast_class = st.selectbox(
                        "Класс давления",
                        options=list(CAST_IRON_BY_CLASS.keys()),
                        index=0,
                        key=f"hyd_cast_class_{mat_code}",
                    )
                    dn_options = [0] + sorted(CAST_IRON_BY_CLASS.get(cast_class, {}).keys())
                    dn_sel = st.selectbox(
                        "DN (номинальный диаметр), мм",
                        options=dn_options,
                        key=f"hyd_dn_{mat_code}",
                    )
            with geom_c3:
                if use_manual_din:
                    st.number_input(
                        "Толщина стенки, мм",
                        min_value=0.0,
                        value=0.0,
                        step=0.1,
                        disabled=True,
                        key=f"hyd_s_disabled_{mat_code}",
                    )
                else:
                    dn_int = int(dn_sel)
                    class_map = CAST_IRON_BY_CLASS.get(cast_class, {})
                    if dn_int > 0 and dn_int in class_map:
                        de_mm, e_min, e_nom = class_map[dn_int]
                        s_mm = float(
                            st.number_input(
                                "Толщина стенки, мм",
                                min_value=float(e_min),
                                max_value=float(e_nom),
                                value=float(e_min),
                                step=0.1,
                                key=f"hyd_s_{mat_code}",
                            )
                        )
                        d_out_mm = float(de_mm)
                        d_in_mm = max(d_out_mm - 2.0 * s_mm, 0.0)
                        st.caption(f"DE = {d_out_mm:.0f} мм (по ГОСТ ISO 2531-2022)")
                    else:
                        st.number_input(
                            "Толщина стенки, мм",
                            min_value=0.0,
                            value=0.0,
                            step=0.1,
                            disabled=True,
                            key=f"hyd_s_{mat_code}",
                        )
        elif mat_code == "fiberglass":
            fg_profiles = FIBERGLASS_DIMENSIONS
            with geom_c2:
                if use_manual_din:
                    st.number_input(
                        "Внутренний диаметр, мм",
                        min_value=0.0,
                        value=0.0,
                        step=0.1,
                        disabled=True,
                        key=f"hyd_din_disabled_{mat_code}",
                    )
                    st.selectbox(
                        "Конструкция",
                        options=list(fg_profiles.keys()),
                        index=0,
                        disabled=True,
                        key=f"hyd_fg_profile_{mat_code}",
                    )
                    st.selectbox(
                        "Рабочее давление, МПа",
                        options=["-"],
                        index=0,
                        disabled=True,
                        key=f"hyd_fg_pressure_{mat_code}",
                    )
                else:
                    fg_profile = st.selectbox(
                        "Конструкция",
                        options=list(fg_profiles.keys()),
                        index=0,
                        key=f"hyd_fg_profile_{mat_code}",
                    )
                    pressure_options = list(fg_profiles.get(fg_profile, {}).keys())
                    fg_pressure = st.selectbox(
                        "Рабочее давление, МПа",
                        options=pressure_options if pressure_options else ["-"],
                        index=0,
                        key=f"hyd_fg_pressure_{mat_code}",
                    )
                    fg_dims_map = fg_profiles.get(fg_profile, {}).get(fg_pressure, {})
            with geom_c3:
                if use_manual_din:
                    st.number_input(
                        "Толщина стенки, мм",
                        min_value=0.0,
                        value=0.0,
                        step=0.1,
                        disabled=True,
                        key=f"hyd_s_disabled_{mat_code}",
                    )
                else:
                    d_in_sel_right = st.selectbox(
                        "Внутренний диаметр, мм",
                        options=[0] + sorted(fg_dims_map.keys()) if fg_dims_map else [0],
                        key=f"hyd_din_right_{mat_code}",
                    )
                    d_in_mm = float(d_in_sel_right)
                    selected_din = int(d_in_mm)
                    if selected_din > 0 and selected_din in fg_dims_map:
                        s_range = fg_dims_map[selected_din]
                        s_min = float(min(s_range))
                        s_max = float(max(s_range))
                        s_mm = float(
                            st.number_input(
                                "Толщина стенки, мм",
                                min_value=s_min,
                                max_value=s_max,
                                value=s_min,
                                step=0.1,
                                key=f"hyd_s_{mat_code}",
                            )
                        )
                        d_out_mm = d_in_mm + 2.0 * s_mm
                    else:
                        st.number_input(
                            "Толщина стенки, мм",
                            min_value=0.0,
                            value=0.0,
                            step=0.1,
                            disabled=True,
                            key=f"hyd_s_{mat_code}",
                        )
        is_steel_material = mat_code in ("steel_vgp", "steel_welded")
        is_steel_or_copper_material = mat_code in ("steel_vgp", "steel_welded", "copper")
        is_fiberglass_material = (mat_code == "fiberglass")
        is_mlpex_or_polyplastic = mat_code in ("metal_plastic", "polyplastic")
        length_in_extra_row = mat_code not in ("plastic", "fiberglass", "steel_vgp", "steel_welded", "cast_iron", "metal_plastic", "polyplastic", "copper")
        clear_in_extra_row = (mat_code in ("plastic", "cast_iron") or is_steel_or_copper_material or is_fiberglass_material or is_mlpex_or_polyplastic)
        with geom_c4:
            if mat_code in ("plastic", "fiberglass", "steel_vgp", "steel_welded", "cast_iron", "metal_plastic", "polyplastic", "copper"):
                hyd_len_m = st.number_input("Длина участка L, м", min_value=0.0, value=len_default, step=1.0, key=len_key)
                if mat_code in ("cast_iron", "fiberglass"):
                    st.markdown("<div style='height: 1.8rem;'></div>", unsafe_allow_html=True)
                    st.button(
                        "Очистить форму",
                        use_container_width=True,
                        key=f"hyd_clear_{mat_code}",
                        on_click=_reset_hyd_form,
                        args=(mat_code,),
                    )
            else:
                st.markdown("<div style='height: 1.7rem;'></div>", unsafe_allow_html=True)
            if not clear_in_extra_row:
                st.button(
                    "Очистить форму",
                    use_container_width=True,
                    key=f"hyd_clear_{mat_code}",
                    on_click=_reset_hyd_form,
                    args=(mat_code,),
                )

        if mat_code == "cast_iron":
            extra_c1, extra_c2, extra_c3, extra_c4 = st.columns([1, 1, 1, 1])
        elif is_steel_or_copper_material:
            extra_c1, extra_c2, extra_c3, extra_c4 = st.columns([1, 1, 1, 1])
        elif is_fiberglass_material:
            extra_c1, extra_c2, extra_c3, extra_c4 = st.columns([1, 1, 1, 1])
        elif mat_code == "plastic":
            extra_c1, extra_c2, extra_c3, extra_c4 = st.columns([1, 1, 1, 1])
        elif is_mlpex_or_polyplastic:
            extra_c1, extra_c2, extra_c3, extra_c4 = st.columns([1, 1, 1, 1])
        else:
            extra_c1, extra_c2 = st.columns([1, 1])
            extra_c3 = None
            extra_c4 = None
        with extra_c1:
            if use_manual_din:
                st.session_state[manual_hint_key] = False
                d_in_mm = float(
                    st.number_input(
                        "Ручной dвн, мм",
                        min_value=0.0,
                        value=manual_default_val,
                        step=1.0,
                        key=manual_value_key,
                    )
                )
            else:
                st.caption("Ручной ввод выключен")
        with extra_c2:
            if length_in_extra_row:
                if mat_code == "cast_iron":
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                elif is_steel_or_copper_material:
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                elif is_fiberglass_material:
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                elif extra_c3 is not None:
                    st.selectbox("Материал PE", options=PLASTIC_PE_GRADES, key=f"hyd_pe_grade_{mat_code}", disabled=use_manual_din)
                else:
                    hyd_len_m = st.number_input("Длина участка L, м", min_value=0.0, value=len_default, step=1.0, key=len_key)
            else:
                st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
        if extra_c3 is not None:
            with extra_c3:
                if mat_code == "cast_iron":
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                elif is_steel_or_copper_material:
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                elif is_fiberglass_material:
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                elif mat_code == "plastic":
                    st.selectbox("Материал PE", options=PLASTIC_PE_GRADES, key=f"hyd_pe_grade_{mat_code}", disabled=use_manual_din)
                elif is_mlpex_or_polyplastic:
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                else:
                    st.markdown("<div style='height: 1.7rem;'></div>", unsafe_allow_html=True)
                    st.button(
                        "Очистить форму",
                        use_container_width=True,
                        key=f"hyd_clear_{mat_code}",
                        on_click=_reset_hyd_form,
                        args=(mat_code,),
                    )
        if extra_c4 is not None:
            with extra_c4:
                if is_steel_or_copper_material:
                    st.button(
                        "Очистить форму",
                        use_container_width=True,
                        key=f"hyd_clear_{mat_code}",
                        on_click=_reset_hyd_form,
                        args=(mat_code,),
                    )
                elif mat_code == "cast_iron":
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                elif is_fiberglass_material:
                    st.markdown("<div style='height: 2.6rem;'></div>", unsafe_allow_html=True)
                else:
                    st.markdown("<div style='height: 1.7rem;'></div>", unsafe_allow_html=True)
                    st.button(
                        "Очистить форму",
                        use_container_width=True,
                        key=f"hyd_clear_{mat_code}",
                        on_click=_reset_hyd_form,
                        args=(mat_code,),
                    )

        if st.session_state.get(manual_hint_key):
            st.info("Найденный диаметр сохранен в ручном вводе. Включите флажок, чтобы применить.")
        st.caption(f"Расчетный dвн = {d_in_mm:.1f} мм")

        c_l1, c_l2, c_l3 = st.columns([1.0, 1.45, 1.55])
        with c_l1:
            local_mode_label = st.selectbox(
                "Учет местных сопротивлений",
                options=["Не учитывать", "По коэффициенту k (СП 30, пп. 8.28-8.29)", "По сумме коэффициентов Σξ"],
                key=f"hyd_local_mode_{mat_code}",
            )
        with c_l2:
            k_mode = "none"
            k_local = 0.0
            k_preset = ""
            xi_sum = 0.0
            if local_mode_label.startswith("По коэффициенту k"):
                k_mode = "k"
                k_preset = st.selectbox("Тип сети (k)", options=list(K_PRESETS.keys()), key=f"hyd_k_preset_{mat_code}")
                k_val = K_PRESETS[k_preset]
                if k_val is None:
                    k_local = float(st.number_input("k (вручную)", min_value=0.0, value=0.3, step=0.05, key=f"hyd_k_manual_{mat_code}"))
                else:
                    k_local = float(k_val)
                    st.caption(f"k = {k_local:.2f}")
            elif local_mode_label.startswith("По сумме"):
                k_mode = "xi"
                xi_sum = float(st.number_input("Σξ", min_value=0.0, value=2.0, step=0.1, key=f"hyd_xi_{mat_code}"))
        with c_l3:
            formula_slot = st.empty()

        if float(d_in_mm) <= 0.0:
            st.warning("Для расчета задайте внутренний диаметр больше 0 мм.")
            hyd_res = calc_hydraulics(
                material=mat_code,
                q_l_s=0.0,
                dp_m=0.01,
                length_m=float(hyd_len_m),
                temp_c=float(hyd_temp_c),
                is_new=bool(hyd_is_new),
                local_mode=k_mode,
                k_local=float(k_local),
                xi_sum=float(xi_sum),
            )
        else:
            hyd_res = calc_hydraulics(
                material=mat_code,
                q_l_s=float(hyd_q_l_s),
                dp_m=float(d_in_mm) / 1000.0,
                length_m=float(hyd_len_m),
                temp_c=float(hyd_temp_c),
                is_new=bool(hyd_is_new),
                local_mode=k_mode,
                k_local=float(k_local),
                xi_sum=float(xi_sum),
            )

        formula_line = ""
        if k_mode == "k":
            formula_line = (
                f"H = i·L·(1+k) = {hyd_res.i_m_per_m:.6f}·{float(hyd_len_m):.2f}·(1+{float(k_local):.2f}) = {hyd_res.h_total_m:.3f} м"
            )
        elif k_mode == "xi":
            formula_line = (
                f"H = i·L + Σξ·v²/(2g) = {hyd_res.h_friction_m:.3f} + {hyd_res.h_local_m:.3f} = {hyd_res.h_total_m:.3f} м"
            )
        else:
            formula_line = (
                f"H = i·L = {hyd_res.i_m_per_m:.6f}·{float(hyd_len_m):.2f} = {hyd_res.h_total_m:.3f} м"
            )
        formula_slot.markdown(
            f"Расчет потерь:  \n**{formula_line}**",
            unsafe_allow_html=False,
        )

        r1, r2, r3, r4, r5 = st.columns(5)
        r1.metric("Скорость v, м/с", f"{hyd_res.v_m_s:.3f}")
        r2.metric("Уклон 1000i, мм/м", f"{hyd_res.i_m_per_m*1000.0:.3f}")
        r3.metric("Потери по длине, м", f"{hyd_res.h_friction_m:.3f}")
        r4.metric("Местные потери, м", f"{hyd_res.h_local_m:.3f}")
        r5.metric("Итого потери H, м", f"{hyd_res.h_total_m:.3f}")

        if hyd_res.v_m_s > v_limit:
            st.error(f"Нарушение скоростного режима: v > {v_limit:.1f} м/с. Рекомендуется увеличить диаметр.")
        else:
            st.success("Скоростной режим в допустимом диапазоне.")
        if hyd_system == "ГВС":
            st.session_state["hyd_last_gvs_total_loss_m"] = float(hyd_res.h_total_m)
            st.session_state["hyd_last_gvs_material"] = str(MATERIALS.get(mat_code, {}).get("label", mat_code))

        cap_col, help_col = st.columns([5, 2])
        with cap_col:
            st.caption(
                f"Re={hyd_res.re:.0f}, λ={hyd_res.lambda_f:.5f}, ν={hyd_res.nu_m2_s:.2e} м²/с, dвн={hyd_res.dp_m*1000:.1f} мм."
            )
        with help_col:
            st.caption("ℹ️ Что такое Re/λ/ν/i")
            if st.checkbox("Показать пояснения", key=f"hyd_help_toggle_{mat_code}"):
                st.markdown(
                    """
- `Re` (число Рейнольдса): показывает режим течения воды (условно ламинарный/турбулентный).
- `ν` (кинематическая вязкость, м²/с): свойство воды, зависит от температуры.
- `λ` (коэффициент трения): насколько труба «тормозит» поток по длине.
- `i` (гидравлический уклон, м/м): потери напора на 1 метр трубы.

Базовые связи:
- `Re = v·d/ν`
- `i = λ·v²/(2gd)`
- `h_f = i·L`
"""
                )

        hyd_png = _build_hydraulic_sheet_png(
            material_label=MATERIALS.get(mat_code, {}).get("label", mat_code),
            system_label=hyd_system,
            scope_label=hyd_scope,
            speed_mode_label=hyd_speed_label,
            q_l_s=float(hyd_q_l_s),
            temp_c=float(hyd_temp_c),
            d_in_mm=float(d_in_mm),
            length_m=float(hyd_len_m),
            local_mode_label=local_mode_label,
            k_local=float(k_local),
            xi_sum=float(xi_sum),
            v_limit=float(v_limit),
            v_m_s=float(hyd_res.v_m_s),
            i_m_per_m=float(hyd_res.i_m_per_m),
            h_friction_m=float(hyd_res.h_friction_m),
            h_local_m=float(hyd_res.h_local_m),
            h_total_m=float(hyd_res.h_total_m),
            re_value=float(hyd_res.re),
            lambda_f=float(hyd_res.lambda_f),
            nu_m2_s=float(hyd_res.nu_m2_s),
        )
        hyd_docx = _build_hydraulic_sheet_docx(
            material_label=MATERIALS.get(mat_code, {}).get("label", mat_code),
            system_label=hyd_system,
            scope_label=hyd_scope,
            speed_mode_label=hyd_speed_label,
            q_l_s=float(hyd_q_l_s),
            temp_c=float(hyd_temp_c),
            d_in_mm=float(d_in_mm),
            length_m=float(hyd_len_m),
            local_mode_label=local_mode_label,
            k_local=float(k_local),
            xi_sum=float(xi_sum),
            v_limit=float(v_limit),
            v_m_s=float(hyd_res.v_m_s),
            i_m_per_m=float(hyd_res.i_m_per_m),
            h_friction_m=float(hyd_res.h_friction_m),
            h_local_m=float(hyd_res.h_local_m),
            h_total_m=float(hyd_res.h_total_m),
            re_value=float(hyd_res.re),
            lambda_f=float(hyd_res.lambda_f),
            nu_m2_s=float(hyd_res.nu_m2_s),
        )
        exp_col_png, exp_col_docx = st.columns(2)
        with exp_col_png:
            _file_export_widget(
                label="⬇️ Экспорт листа гидравлики (PNG)",
                data=hyd_png,
                file_name=f"hydraulic_sheet_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                key=f"hyd_sheet_png_{mat_code}",
                mime="image/png",
            )
        with exp_col_docx:
            _doc_export_widget(
                label="⬇️ Экспорт листа гидравлики (Word)",
                data=hyd_docx,
                file_name=f"hydraulic_sheet_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                key=f"hyd_sheet_docx_{mat_code}",
            )

        st.markdown("### Расчет требуемого напора и счетчиков")
        st.caption("СП 30.13330.2020: п. 8.27 и п. 12")
        st.markdown("`Hтр = Hгеом + Hпотерь + Hсвоб + hсч + Hтепл + Hввод`")

        meter_table = [
            {"dn": 15, "q_exp_m3_h": 1.2, "q_max_m3_h": 3.0, "s": 14.5},
            {"dn": 20, "q_exp_m3_h": 2.0, "q_max_m3_h": 5.0, "s": 5.18},
            {"dn": 25, "q_exp_m3_h": 2.8, "q_max_m3_h": 7.0, "s": 2.64},
            {"dn": 32, "q_exp_m3_h": 4.0, "q_max_m3_h": 10.0, "s": 1.3},
            {"dn": 40, "q_exp_m3_h": 6.4, "q_max_m3_h": 16.0, "s": 0.5},
            {"dn": 50, "q_exp_m3_h": 12.0, "q_max_m3_h": 30.0, "s": 0.143},
            {"dn": 65, "q_exp_m3_h": 17.0, "q_max_m3_h": 70.0, "s": 810e-5},
            {"dn": 80, "q_exp_m3_h": 36.0, "q_max_m3_h": 110.0, "s": 264e-5},
            {"dn": 100, "q_exp_m3_h": 65.0, "q_max_m3_h": 180.0, "s": 76.6e-5},
            {"dn": 150, "q_exp_m3_h": 140.0, "q_max_m3_h": 350.0, "s": 13e-5},
            {"dn": 200, "q_exp_m3_h": 210.0, "q_max_m3_h": 600.0, "s": 3.5e-5},
            {"dn": 250, "q_exp_m3_h": 380.0, "q_max_m3_h": 1000.0, "s": 1.8e-5},
        ]

        def _meter_type_by_dn(dn: int) -> str:
            return "Крыльчатый" if int(dn) <= 40 else "Турбинный"

        def _meter_limit_m(dn: int, with_fire: bool) -> float:
            meter_type = _meter_type_by_dn(dn)
            if with_fire:
                return 10.0 if meter_type == "Крыльчатый" else 5.0
            return 5.0 if meter_type == "Крыльчатый" else 2.5

        def _pick_meter(q_avg_m3_h: float, q_l_s: float, with_fire: bool) -> dict:
            qavg = max(float(q_avg_m3_h), 0.0)
            qs = max(float(q_l_s), 0.0)
            # СП 30.13330.2020 п.12.14: предварительный подбор по среднечасовому расходу
            # с проверкой на эксплуатационный расход из табл. 12.1.
            by_avg = [r for r in meter_table if qavg <= float(r["q_exp_m3_h"])]
            pool = by_avg if by_avg else [meter_table[-1]]
            chosen = None
            for r in pool:
                dn = int(r["dn"])
                s = float(r["s"])
                h = s * (qs ** 2)
                lim = _meter_limit_m(dn, with_fire)
                if h <= lim:
                    chosen = dict(r)
                    break
            if chosen is None:
                chosen = dict(pool[-1])
            dn = int(chosen["dn"])
            s = float(chosen["s"])
            h = s * (qs ** 2)
            lim = _meter_limit_m(dn, with_fire)
            return {
                "dn": dn,
                "q_avg_m3_h": qavg,
                "q_l_s": qs,
                "q_exp_m3_h": float(chosen["q_exp_m3_h"]),
                "q_max_m3_h": float(chosen["q_max_m3_h"]),
                "s": s,
                "h_m": h,
                "limit_m": lim,
                "ok": h <= lim,
                "meter_type": _meter_type_by_dn(dn),
            }

        h_geo_default = float(st.session_state.get("passport_h_top", 0.0) or 0.0)
        h_free_default = float(st.session_state.get("passport_free_head_m", 20.0) or 20.0)
        q_hvs_default = float(water_res.get("cold_max_l_sec", 0.0)) if "water_res" in locals() else 0.0
        q_gvs_default = float(water_res.get("hot_max_l_sec", 0.0)) if "water_res" in locals() else 0.0
        qavg_hvs_default = float(water_res.get("cold_avg_m3_hour", 0.0)) if "water_res" in locals() else 0.0
        qavg_gvs_default = float(water_res.get("hot_avg_m3_hour", 0.0)) if "water_res" in locals() else 0.0
        _sync_autofill_number(f"hyd_head_geo_{mat_code}", h_geo_default)
        _sync_autofill_number(f"hyd_head_free_{mat_code}", h_free_default)

        n1, n2, n3, n4, n5 = st.columns(5)
        with n1:
            h_geo_m = float(st.number_input("Hгеом, м", min_value=0.0, value=h_geo_default, step=0.1, key=f"hyd_head_geo_{mat_code}"))
        with n2:
            h_losses_m = float(
                st.number_input(
                    "Hпотерь (гидравлика), м",
                    min_value=0.0,
                    value=float(hyd_res.h_total_m),
                    step=0.001,
                    key=f"hyd_head_losses_{mat_code}",
                    disabled=True,
                )
            )
        with n3:
            h_free_m = float(st.number_input("Hсвоб, м", min_value=0.0, value=h_free_default, step=0.1, key=f"hyd_head_free_{mat_code}"))
        with n4:
            l_inlet_m = float(
                st.number_input(
                    "Lввода, м",
                    min_value=0.0,
                    value=0.0,
                    step=0.1,
                    key=f"hyd_inlet_len_{mat_code}",
                    help="Отдельная величина: длина ввода, включая вертикальную часть.",
                )
            )
        with n5:
            st.caption("Hпотерь и уклон i подтягиваются из текущего расчета гидравлики.")

        st.markdown("**Потери в счетчиках (п. 12): `hсч = S · q²`**")
        st.caption("Подбор DN: по п.12.14 (`qсрч` против `Qэкспл` табл. 12.1). Проверка потерь: по п.12.15/12.16.")
        f1, f2, f3, f4 = st.columns([1.45, 0.9, 1.35, 0.9])
        with f1:
            fire_mode = st.checkbox(
                "Проверка счетчика с пожарным расходом (п. 12.16б)",
                value=False,
                key=f"hyd_meter_fire_{mat_code}",
            )
        q_fire_meter_l_s = 0.0
        with f2:
            if fire_mode:
                q_fire_meter_l_s = float(
                    st.number_input(
                        "qпож, л/с",
                        min_value=0.0,
                        value=0.0,
                        step=0.001,
                        key=f"hyd_meter_q_fire_{mat_code}",
                    )
                )
        with f3:
            has_fire_pipeline = st.checkbox(
                "Наличие противопожарного водопровода",
                value=False,
                key=f"hyd_has_fire_pipeline_{mat_code}",
            )
            st.caption("СП 10.13130.2020: табл. 7.3, п. 7.15")
        h_fire_valve_m = 0.0
        with f4:
            if has_fire_pipeline:
                h_fire_valve_m = float(
                    st.number_input(
                        "Hпож, м",
                        min_value=0.0,
                        value=0.0,
                        step=0.1,
                        key=f"hyd_fire_head_m_{mat_code}",
                    )
                )
        q1, q2, q3 = st.columns([1.0, 1.0, 1.5])
        if hyd_system == "ГВС":
            _sync_autofill_number(f"hyd_meter_qh_gvs_{mat_code}", qavg_gvs_default)
            _sync_autofill_number(f"hyd_meter_q_gvs_{mat_code}", q_gvs_default)
            qh_active_default = qavg_gvs_default
            q_active_default = q_gvs_default
            with q1:
                qh_active = float(
                    st.number_input(
                        "qсрч ГВС, м³/ч",
                        min_value=0.0,
                        value=qh_active_default,
                        step=0.001,
                        key=f"hyd_meter_qh_gvs_{mat_code}",
                    )
                )
            with q2:
                q_active = float(
                    st.number_input(
                        "qрасч,max ГВС, л/с",
                        min_value=0.0,
                        value=q_active_default,
                        step=0.001,
                        key=f"hyd_meter_q_gvs_{mat_code}",
                    )
                )
        else:
            _sync_autofill_number(f"hyd_meter_qh_hvs_{mat_code}", qavg_hvs_default)
            _sync_autofill_number(f"hyd_meter_q_hvs_{mat_code}", q_hvs_default)
            qh_active_default = qavg_hvs_default
            q_active_default = q_hvs_default
            with q1:
                qh_active = float(
                    st.number_input(
                        "qсрч ХВС, м³/ч",
                        min_value=0.0,
                        value=qh_active_default,
                        step=0.001,
                        key=f"hyd_meter_qh_hvs_{mat_code}",
                    )
                )
            with q2:
                q_active = float(
                    st.number_input(
                        "qрасч,max ХВС, л/с",
                        min_value=0.0,
                        value=q_active_default,
                        step=0.001,
                        key=f"hyd_meter_q_hvs_{mat_code}",
                    )
                )
        with q3:
            st.caption("Расходы автоподхватываются из общего расчета воды, при необходимости можно изменить вручную.")

        q_meter_check_l_s = float(q_active) + (float(q_fire_meter_l_s) if fire_mode else 0.0)
        meter_active = _pick_meter(qh_active, q_meter_check_l_s, fire_mode)
        has_itp_heating = bool(st.session_state.get("passport_has_itp_heating", False))
        h_hex_m = 3.0 if has_itp_heating else 0.0
        # В Hтр учитываем потери счетчика на хозяйственный расход;
        # пожарный расход используется только для проверки счетчика по п.12.16б.
        h_meter_active_m = float(meter_active["s"]) * (float(q_active) ** 2)
        h_inlet_m = float(hyd_res.i_m_per_m) * float(l_inlet_m)
        h_free_for_sum_m = max(float(h_free_m), float(h_fire_valve_m))
        h_required_m = h_geo_m + h_losses_m + h_free_for_sum_m + h_meter_active_m + h_hex_m + h_inlet_m
        st.session_state["hyd_required_head_m"] = float(h_required_m)
        st.session_state["hyd_required_pressure_mpa"] = float(h_required_m) * 0.00980665
        st.session_state["hyd_inlet_loss_m"] = float(h_inlet_m)
        if hyd_system == "ГВС":
            # Паспорт ГВС берет эти значения автоматически из гидравлики.
            st.session_state["passport_h_top"] = float(h_geo_m)
            st.session_state["passport_free_head_m"] = float(h_free_m)
            st.session_state["passport_meter_loss_m"] = float(h_meter_active_m)
            st.session_state["passport_has_meter"] = True
            st.session_state["hyd_required_head_m_gvs"] = float(h_required_m)
        elif hyd_system == "ХВС":
            st.session_state["hyd_required_head_m_hvs"] = float(h_required_m)

        m1, m2, m3 = st.columns([1.4, 1.0, 1.0])
        with m1:
            st.caption(
                f"{hyd_system}: DN {meter_active['dn']} ({meter_active['meter_type']}), "
                f"Qэкспл={meter_active['q_exp_m3_h']:.1f} м³/ч, "
                f"S={meter_active['s']:.5f}, hсч.проверка={meter_active['h_m']:.3f} м, допуск {meter_active['limit_m']:.1f} м"
            )
            if bool(meter_active["ok"]):
                st.caption("Потери в счетчике в допустимом диапазоне.")
            else:
                st.caption("Потери в счетчике выше допустимых — нужен счетчик большего DN/типа.")
            if fire_mode:
                st.caption(
                    f"Проверка счетчика: qпроверки = qрасч,max + qпож = "
                    f"{q_active:.3f} + {q_fire_meter_l_s:.3f} = {q_meter_check_l_s:.3f} л/с"
                )
            st.caption(f"Hтепл (ИТП): {h_hex_m:.1f} м")
            st.caption(f"Hввод = i·Lввода = {hyd_res.i_m_per_m:.6f}·{l_inlet_m:.2f} = {h_inlet_m:.3f} м")
            if has_fire_pipeline:
                st.caption(
                    f"Hсвоб,прин = max(Hсвоб={h_free_m:.3f}; Hпож={h_fire_valve_m:.3f}) = {h_free_for_sum_m:.3f} м"
                )
        with m2:
            st.metric("hсч (текущая система), м", f"{h_meter_active_m:.3f}")
        with m3:
            st.metric("Требуемый напор Hтр, м", f"{h_required_m:.3f}")
            st.caption(f"Система: {hyd_system}")
        st.caption(
            "Hтр = "
            f"{h_geo_m:.3f} (Hгеом) + "
            f"{h_losses_m:.3f} (Hпотерь) + "
            f"{h_free_for_sum_m:.3f} (Hсвоб,прин) + "
            f"{h_meter_active_m:.3f} (hсч) + "
            f"{h_hex_m:.3f} (Hтепл) + "
            f"{h_inlet_m:.3f} (Hввод) = "
            f"{h_required_m:.3f} м"
        )
        head_meter_docx = _build_head_meter_calc_docx(
            material_label=MATERIALS.get(mat_code, {}).get("label", mat_code),
            scope_label=hyd_scope,
            system_label=hyd_system,
            h_geo_m=h_geo_m,
            h_losses_m=h_losses_m,
            h_free_m=h_free_for_sum_m,
            h_hex_m=h_hex_m,
            h_inlet_m=h_inlet_m,
            i_m_per_m=float(hyd_res.i_m_per_m),
            l_inlet_m=float(l_inlet_m),
            meter_active=meter_active,
            h_meter_active_m=h_meter_active_m,
            h_required_m=h_required_m,
            fire_mode=bool(fire_mode),
        )
        hm_btn1, hm_btn2 = st.columns(2)
        with hm_btn1:
            st.button(
                "Очистить форму (напор и счетчики)",
                use_container_width=True,
                key=f"hyd_head_meter_clear_{mat_code}",
                on_click=_reset_hyd_head_meter_form,
                args=(mat_code,),
            )
        with hm_btn2:
            _doc_export_widget(
                label="⬇️ Word: расчет напора и счетчиков",
                data=head_meter_docx,
                file_name=f"hydraulic_head_meter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                key=f"hyd_head_meter_docx_{mat_code}",
            )

with tab_gvs:
    st.subheader("Параметры паспорта ГВС")
    st.caption("Расчеты паспорта ГВС по СП 30.13330.2020.")
    annex_gvs = str(st.session_state.get("annex_gvs", "") or "").strip()
    fixture_q0_map = {
        "Раковина": {"q0_l_s": 0.07, "q0hr_l_h": 40.0},
        "Мойка": {"q0_l_s": 0.14, "q0hr_l_h": 60.0},
        "Мойка двойная": {"q0_l_s": 0.20, "q0hr_l_h": 100.0},
        "Душ": {"q0_l_s": 0.14, "q0hr_l_h": 60.0},
        "Поддон": {"q0_l_s": 0.14, "q0hr_l_h": 60.0},
        "Ванна": {"q0_l_s": 0.18, "q0hr_l_h": 60.0},
        "Биде": {"q0_l_s": 0.10, "q0hr_l_h": 40.0},
    }
    fixture_key_map = {
        "Раковина": "passport_f_sink",
        "Мойка": "passport_f_moyka",
        "Мойка двойная": "passport_f_double",
        "Душ": "passport_f_shower",
        "Поддон": "passport_f_tray",
        "Ванна": "passport_f_bath",
        "Биде": "passport_f_bidet",
    }
    g1, g2, g3, g4 = st.columns(4)
    with g1:
        t_hot_c = st.number_input("Температура горячей воды th, °C", min_value=45.0, max_value=75.0, value=60.0, step=1.0, key="t_hot_c")
    with g2:
        t_cold_c = st.number_input("Температура холодной воды tc, °C", min_value=1.0, max_value=25.0, value=10.0, step=1.0, key="t_cold_c")
    with g3:
        qht_kw = st.number_input("Потери тепла трубопроводами Qht, кВт", min_value=0.0, value=0.0, step=0.1, key="qht_kw")
    with g4:
        delta_t_supply_c = st.number_input("Δt для циркуляции (t3 - t4), °C", min_value=1.0, value=10.0, step=1.0, key="delta_t_supply_c")

    calc_rows_live = _apply_a2_notes_rules(
        rows=st.session_state.water_consumers,
        selected_object_kind=selected_object_kind,
        shift_count=int(shift_count),
        shift_hours=float(shift_hours),
        apply_shift_rules=bool(apply_shift_rules),
        apartment_rooms_k=int(apartment_rooms_k),
        use_apartment_formula=bool(use_apartment_formula),
        use_food_formula=bool(use_food_formula),
        food_seats_n=float(food_seats_n),
        food_m=float(food_m),
        food_t_hours=float(food_t_hours),
        food_y=float(food_y),
        prod_household_coeff=float(prod_household_coeff),
        laundry_hot_uplift_pct=float(laundry_hot_uplift_pct),
        use_global_work_hours=bool(use_global_work_hours),
        global_work_hours=float(global_work_hours),
    )
    water_models = _consumers_to_models(calc_rows_live)
    water_res_live = calc_water_by_consumers_advanced(
        consumers=water_models,
        peak_hour_factor=PEAK_HOUR_FACTOR,
        day_factor=DAY_FACTOR,
        reserve_factor=RESERVE_FACTOR_WATER,
        leakage_percent=LEAKAGE_PERCENT,
    )
    auto_qh_avg_m3_h = float(water_res_live["hot_avg_m3_hour"])
    auto_qh_max_m3_h = float(water_res_live["hot_max_m3_hour"])

    if "gvs_manual_mode" not in st.session_state:
        st.session_state["gvs_manual_mode"] = False
    if "gvs_manual_qh_avg_m3_h" not in st.session_state:
        st.session_state["gvs_manual_qh_avg_m3_h"] = auto_qh_avg_m3_h
    if "gvs_manual_qh_max_m3_h" not in st.session_state:
        st.session_state["gvs_manual_qh_max_m3_h"] = auto_qh_max_m3_h
    if "gvs_manual_kcir" not in st.session_state:
        st.session_state["gvs_manual_kcir"] = 0.3
    if not st.session_state.get("gvs_qht_default_migrated_v1", False):
        if float(st.session_state.get("qht_kw", 0.0) or 0.0) == 2.88:
            st.session_state["qht_kw"] = 0.0
        st.session_state["gvs_qht_default_migrated_v1"] = True

    st.caption(
        f"Автоподхват из 'Баланс ВС и ВО' по умолчанию: qT_h={auto_qh_avg_m3_h:.3f} м3/ч, "
        f"qhr_h={auto_qh_max_m3_h:.3f} м3/ч"
    )
    st.checkbox("Ручной ввод расходов/вероятностей", key="gvs_manual_mode")

    qh_avg_input_m3_h = auto_qh_avg_m3_h
    qh_max_input_m3_h = auto_qh_max_m3_h
    manual_kcir_enabled = False
    if st.session_state.get("gvs_manual_mode", False):
        man1, man2, man3, man4 = st.columns([1, 1, 0.9, 1.1])
        with man1:
            qh_avg_input_m3_h = st.number_input(
                "qT_h (средний), м3/ч",
                min_value=0.0,
                step=0.001,
                key="gvs_manual_qh_avg_m3_h",
            )
        with man2:
            qh_max_input_m3_h = st.number_input(
                "qhr_h (макс.), м3/ч",
                min_value=0.0,
                step=0.001,
                key="gvs_manual_qh_max_m3_h",
            )
        with man3:
            st.markdown('<div style="margin-top: 1.85rem;"></div>', unsafe_allow_html=True)
            manual_kcir_enabled = st.checkbox("kcir вручную", key="gvs_manual_kcir_enabled")
        with man4:
            st.number_input(
                "kcir",
                min_value=0.0,
                max_value=2.0,
                step=0.01,
                key="gvs_manual_kcir",
                disabled=not bool(manual_kcir_enabled),
            )

    gvs_res = calc_gvs_passport(
        qh_avg_m3_h=float(qh_avg_input_m3_h),
        qh_max_m3_h=float(qh_max_input_m3_h),
        t_hot_c=float(t_hot_c),
        t_cold_c=float(t_cold_c),
        qht_kW=float(qht_kw),
        delta_t_supply_c=float(delta_t_supply_c),
    )
    if st.session_state.get("gvs_manual_mode", False) and bool(manual_kcir_enabled):
        qh_l_s = float(gvs_res["qh_l_s"])
        kcir_manual = max(float(st.session_state.get("gvs_manual_kcir", 0.0) or 0.0), 0.0)
        qh_cir_l_s = qh_l_s * (1.0 + kcir_manual)
        gvs_res["kcir"] = kcir_manual
        gvs_res["qh_cir_l_s"] = qh_cir_l_s
        gvs_res["qh_cir_m3_h"] = qh_cir_l_s * 3.6

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("QTh средний, кВт", f'{gvs_res["qth_kW"]:.3f}')
    m2.metric("Qhr,h максимальный, кВт", f'{gvs_res["qhrh_kW"]:.3f}')
    m3.metric("qcir, л/с", f'{gvs_res["qcir_l_s"]:.4f}')
    m4.metric("qh,cir, л/с", f'{gvs_res["qh_cir_l_s"]:.4f}')

    if st.session_state.get("gvs_manual_mode", False):
        st.caption(
            f"Ручной ввод: qT_h={gvs_res['qh_avg_m3_h']:.3f} м3/ч, "
            f"qhr_h={gvs_res['qh_max_m3_h']:.3f} м3/ч, kcir={gvs_res['kcir']:.3f}"
        )
    else:
        st.caption(
            f"Автоподхват расходов из вкладки 'Баланс ВС и ВО': qT_h={gvs_res['qh_avg_m3_h']:.3f} м3/ч, "
            f"qhr_h={gvs_res['qh_max_m3_h']:.3f} м3/ч, kcir={gvs_res['kcir']:.3f}"
        )

    # Паспорт ГВС: ручные поля и выбор характерного прибора (А.1)
    hot_rows = [r for r in water_res_live["rows"] if float(r.get("count", 0.0) or 0.0) > 0 and float(r.get("hot_m3_day", 0.0) or 0.0) > 0]
    auto_hours = float(global_work_hours) if bool(use_global_work_hours) else 24.0
    if not bool(use_global_work_hours) and hot_rows:
        auto_hours = max(float(r.get("t_hours", 24.0) or 24.0) for r in hot_rows)

    with st.expander("Паспорт ГВС: исходные данные", expanded=True):
        passport_consumers_auto = float(
            sum(
                float(r.get("count", 0.0) or 0.0)
                for r in water_res_live["rows"]
                if float(r.get("count", 0.0) or 0.0) > 0 and _is_people_unit(str(r.get("unit", "")))
            )
        )
        auto_seed = {
            "passport_consumers": 0.0,
            "passport_area": 0.0,
            "passport_volume": 0.0,
            "passport_devices_total": 0.0,
            "passport_hours_per_day": 0.0,
            "passport_h_top": 0.0,
            "passport_free_head_m": 20.0,
            "passport_losses_system_m": 0.0,
            "passport_circ_losses_m": 0.0,
            "passport_has_meter": True,
            "passport_has_itp_heating": False,
            "passport_meter_loss_m": 0.0,
        }
        for fk in fixture_key_map.values():
            auto_seed[fk] = 0.0
        for k, v in auto_seed.items():
            if k not in st.session_state:
                st.session_state[k] = v
        # Миграция старого дефолта (0.5 м) на нулевой старт.
        if not st.session_state.get("passport_meter_loss_m_migrated_v2", False):
            if float(st.session_state.get("passport_meter_loss_m", 0.0) or 0.0) == 0.5:
                st.session_state["passport_meter_loss_m"] = 0.0
            st.session_state["passport_meter_loss_m_migrated_v2"] = True
        # Миграция v3: начальные значения потребителей/высоты диктующего прибора должны быть нулевыми.
        if not st.session_state.get("passport_seed_migrated_v3", False):
            if float(st.session_state.get("passport_h_top", 0.0) or 0.0) == 7.1:
                st.session_state["passport_h_top"] = 0.0
            if float(st.session_state.get("passport_consumers", 0.0) or 0.0) == float(passport_consumers_auto):
                st.session_state["passport_consumers"] = 0.0
            if float(st.session_state.get("passport_hours_per_day", 0.0) or 0.0) == float(auto_hours):
                st.session_state["passport_hours_per_day"] = 0.0
            st.session_state["passport_seed_migrated_v3"] = True
        # Миграция v4: свободный напор по умолчанию 20 м.
        if not st.session_state.get("passport_free_head_m_migrated_v4", False):
            if float(st.session_state.get("passport_free_head_m", 0.0) or 0.0) <= 0.0:
                st.session_state["passport_free_head_m"] = 20.0
            st.session_state["passport_free_head_m_migrated_v4"] = True

        b1, b2 = st.columns(2)
        with b1:
            st.info("Расходы и kcir задаются выше: автоподхватом из баланса или ручным вводом.")
        with b2:
            if st.button("Очистить форму паспорта", use_container_width=True):
                for k, v in auto_seed.items():
                    st.session_state[k] = v
                st.success("Форма паспорта очищена.")

        p1, p2, p3, p4 = st.columns(4)
        with p1:
            passport_consumers = st.number_input("2) Количество основных потребителей", min_value=0.0, step=1.0, key="passport_consumers")
        with p2:
            passport_area = st.number_input("3) Общая площадь, м²", min_value=0.0, step=10.0, key="passport_area")
        with p3:
            passport_volume = st.number_input("4) Строительный объем, м³", min_value=0.0, step=10.0, key="passport_volume")
        with p4:
            passport_devices_total = st.number_input("5) Общее кол-во санприборов, шт.", min_value=0.0, step=1.0, key="passport_devices_total")

        hyd_mat = str(st.session_state.get("hyd_material", "") or "")
        hyd_sys = str(st.session_state.get(f"hyd_system_{hyd_mat}", "") or "") if hyd_mat else ""
        hyd_geo_key = f"hyd_head_geo_{hyd_mat}" if hyd_mat else ""
        hyd_free_key = f"hyd_head_free_{hyd_mat}" if hyd_mat else ""
        if hyd_sys == "ГВС" and hyd_geo_key and hyd_geo_key in st.session_state:
            st.session_state["passport_h_top"] = float(st.session_state.get(hyd_geo_key, 0.0) or 0.0)
        if hyd_sys == "ГВС" and hyd_free_key and hyd_free_key in st.session_state:
            st.session_state["passport_free_head_m"] = float(st.session_state.get(hyd_free_key, 20.0) or 20.0)

        pp1, pp2, pp3, pp4 = st.columns(4)
        with pp1:
            passport_hours_per_day = st.number_input("6) Число часов работы в сутки, ч/сут", min_value=0.0, max_value=24.0, step=1.0, key="passport_hours_per_day")
        with pp2:
            fixture_name = st.selectbox(
                "7) Характерный прибор (Приложение А.1 СП 30.13330.2020)",
                options=list(fixture_q0_map.keys()),
                index=0,
            )
        with pp3:
            fixture_q0 = st.number_input(
                "7) Расход характерного прибора, л/с",
                min_value=0.01,
                value=float(fixture_q0_map.get(fixture_name, {}).get("q0_l_s", 0.07)),
                step=0.01,
            )
        with pp4:
            st.markdown('<div style="margin-top: 1.85rem;"></div>', unsafe_allow_html=True)
            passport_has_itp_heating = st.checkbox("Подогрев воды в ИТП", key="passport_has_itp_heating")

        passport_h_top = float(st.session_state.get("passport_h_top", 0.0) or 0.0)
        free_head_m = float(st.session_state.get("passport_free_head_m", 20.0) or 20.0)
        passport_has_meter = bool(st.session_state.get("passport_has_meter", True))
        meter_loss_m = float(st.session_state.get("passport_meter_loss_m", 0.0) or 0.0)
        h_required_auto = float(st.session_state.get("hyd_required_head_m_gvs", 0.0) or 0.0)

        st.caption("Потери давления и потребный напор рассчитываются в разделе «Гидравлика» и автоматически подтягиваются в паспорт ГВС.")
        rr1, rr2, rr3, rr4 = st.columns(4)
        with rr1:
            st.metric("17) Высота дикт. прибора, м", f"{passport_h_top:.2f}")
        with rr2:
            st.metric("Свободный напор, м", f"{free_head_m:.2f}")
        with rr3:
            st.metric("Потери в водомере, м", f"{meter_loss_m:.2f}")
        with rr4:
            st.metric("Hтр (из гидравлики), м", f"{h_required_auto:.2f}")

        hyd_gvs_losses_m = float(st.session_state.get("hyd_last_gvs_total_loss_m", 0.0) or 0.0)
        if hyd_gvs_losses_m > 0.0:
            st.session_state["passport_losses_system_m"] = hyd_gvs_losses_m
        losses_system_m = float(st.session_state.get("passport_losses_system_m", 0.0) or 0.0)
        circ_losses_m = float(st.session_state.get("passport_circ_losses_m", 0.0) or 0.0)

        st.caption("Состав приборов (для печати в паспорте):")
        fcols = st.columns(len(fixture_key_map))
        for col, (fname, fkey) in zip(fcols, fixture_key_map.items()):
            with col:
                st.number_input(f"{fname}, шт.", min_value=0.0, step=1.0, key=fkey)

    object_designation = object_name
    if water_res_live["rows"]:
        ranked = sorted(
            [r for r in water_res_live["rows"] if float(r.get("count", 0.0) or 0.0) > 0],
            key=lambda x: float(x.get("count", 0.0) or 0.0),
            reverse=True,
        )
        if ranked:
            object_designation = str(ranked[0].get("name", object_name))

    passport_inputs = {
        "consumers_count": float(passport_consumers),
        "area_m2": float(passport_area),
        "volume_m3": float(passport_volume),
        "devices_total": float(passport_devices_total),
        "hours_per_day": float(passport_hours_per_day),
        "fixture_name": fixture_name,
        "q0_char_l_s": float(fixture_q0),
        "q0hr_char_l_h": float(fixture_q0_map.get(fixture_name, {}).get("q0hr_l_h", 60.0)),
        "h_top_m": float(passport_h_top),
        "free_head_m": float(free_head_m),
        "losses_system_m": float(losses_system_m),
        "circ_losses_m": float(circ_losses_m),
        "has_meter": bool(passport_has_meter),
        "has_itp_heating": bool(passport_has_itp_heating),
        "meter_loss_m": float(meter_loss_m),
        "fixture_counts": {fname: float(st.session_state.get(fkey, 0.0) or 0.0) for fname, fkey in fixture_key_map.items()},
    }

    gvs_doc = build_gvs_passport_docx(
        object_name=object_designation,
        water_rows=water_res_live["rows"],
        gvs=gvs_res,
        passport_inputs=passport_inputs,
        annex_label=(annex_gvs or "").strip(),
    )
    _doc_export_widget(
        label="⬇️ Скачать Word-паспорт ГВС",
        data=gvs_doc,
        file_name="passport_gvs.docx",
        key="gvs_passport",
    )
