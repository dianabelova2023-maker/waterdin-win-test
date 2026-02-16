from __future__ import annotations

from io import BytesIO
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.2
    style.paragraph_format.space_after = Pt(0)


def _add_kv_table(doc: Document, rows: List[Tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)


def _is_production_object(consumers: List[Dict[str, float | str]]) -> bool:
    for c in consumers:
        if float(c.get("count", 0.0) or 0.0) <= 0:
            continue
        if str(c.get("object_kind", "nonproduction")).strip().lower() == "production":
            return True
    return False


def _fmt2(value: float) -> str:
    return f"{float(value or 0.0):.2f}"


def _fmt_or_blank(value: float, eps: float = 1e-9) -> str:
    v = float(value or 0.0)
    if abs(v) <= eps:
        return ""
    return f"{v:.2f}"


def _fmt_or_blank_prec(value: float, digits: int, eps: float = 1e-12) -> str:
    v = float(value or 0.0)
    if abs(v) <= eps:
        return ""
    return f"{v:.{max(int(digits), 0)}f}"


def _optf(opts: Dict[str, str], key: str, default: float = 0.0) -> float:
    try:
        return float(opts.get(key, default) or default)
    except Exception:
        return default


def _set_cell_center(cell) -> None:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_vertical(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    text_direction = OxmlElement("w:textDirection")
    text_direction.set(qn("w:val"), "btLr")
    tc_pr.append(text_direction)


def _set_cell_horizontal(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:textDirection"):
            tc_pr.remove(child)
    _set_cell_center(cell)


def _set_table_font_size(table, size_pt: int = 9) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.size = Pt(size_pt)


def _set_cell_text_center(cell, text: str, bold: bool = False) -> None:
    cell.text = str(text)
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].bold = bold
    _set_cell_center(cell)


def _set_col_widths(table, widths_cm: List[float]) -> None:
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _is_irrigation_name(name: str) -> bool:
    n = (name or "").strip().lower()
    return ("полив" in n) or ("заливка поверхности катка" in n)


def _is_no_time_consumer(name: str) -> bool:
    n = (name or "").strip().lower()
    markers = ["полив", "заливка поверхности катка", "подпитка котельной"]
    return any(m in n for m in markers)


def _infer_prod_water_quality(name: str, override: str) -> str:
    ov = (override or "").strip()
    if ov:
        return ov
    n = (name or "").strip().lower()
    if "оборот" in n:
        return "Оборотная"
    tech_markers = [
        "цех",
        "производ",
        "завод",
        "лаборатор",
        "прачеч",
        "котельн",
        "подпитка котельной",
    ]
    if any(m in n for m in tech_markers):
        return "Техническая"
    return "Питьевая (СанПиН 2.1.3684-21, 1.2.3685-21)"


def _infer_consumption_mode(name: str, t_hours: float) -> str:
    n = (name or "").strip().lower()
    if ("заливка поверхности катка" in n) or ("слив" in n and "бассейн" in n):
        return "Залповый"
    periodic_markers = ["душ", "полив", "бан", "бассейн", "мойка", "прачеч"]
    if any(m in n for m in periodic_markers):
        return "Периодический"
    shift_markers = ["цех", "производ", "завод", "рабоч", "предприят", "лаборатор"]
    if any(m in n for m in shift_markers) or (0 < float(t_hours or 0.0) <= 16):
        return "По сменам"
    if float(t_hours or 0.0) >= 20:
        return "Постоянный"
    return "Периодический"


def _infer_sewer_characteristic(name: str, sewer_target: str, water_quality: str) -> str:
    n = (name or "").strip().lower()
    wq = (water_quality or "").strip().lower()
    route = (sewer_target or "").strip().lower()

    if "конденсат" in n:
        return "Конденсат"
    if "дренаж" in n:
        return "Дренажные воды"
    if "оборот" in n or "условно" in n or "оборот" in wq:
        return "Условно-чистые"

    organic_markers = ["предприятия общественного питания", "столов", "пищеблок", "ресторан", "кафе", "буфет"]
    mech_markers = ["цех", "завод", "производ", "прачеч", "лаборатор", "мойк", "бассейн"]
    domestic_markers = ["туалет", "сануз", "душ", "ванн", "раковин", "умываль", "жиль", "санатор", "гостиниц", "общежит"]

    if any(m in n for m in organic_markers):
        return "Производственные (органические примеси)"
    if any(m in n for m in mech_markers):
        return "Производственные (механические примеси)"
    if any(m in n for m in domestic_markers):
        return "Бытовые"
    if route == "production":
        return "Производственные"
    return "Бытовые"


def _infer_form2_sewer_col(name: str) -> int:
    n = (name or "").strip().lower()
    organic_markers = ["предприятия общественного питания", "столов", "пищеблок", "ресторан", "кафе", "буфет"]
    mech_markers = ["цех", "завод", "производ", "прачеч", "лаборатор", "бассейн", "душевые в бытовых помещениях промышленных предприятий"]
    severe_med_markers = ["инфекцион", "патолого", "морг", "кров"]
    if any(m in n for m in organic_markers):
        return 17
    if any(m in n for m in mech_markers):
        return 16
    if any(m in n for m in severe_med_markers):
        return 17
    return 14


def _add_form2_balance_table(
    doc: Document,
    object_name: str,
    object_address: str,
    consumers: List[Dict[str, float | str]],
    water_results: Dict[str, float],
    water_options: Dict[str, str],
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("БАЛАНС ВОДОСНАБЖЕНИЯ И ВОДООТВЕДЕНИЯ").bold = True
    obj_line = object_name.strip() if object_name else "______________________________"
    addr_line = object_address.strip() if object_address else "______________________________"
    doc.add_paragraph(f"Наименование объекта: {obj_line}")
    doc.add_paragraph(f"Адрес объекта: {addr_line}")

    table = doc.add_table(rows=4, cols=19)
    table.style = "Table Grid"
    _set_col_widths(
        table,
        [0.6, 2.8, 1.3, 1.0, 1.2, 1.5, 1.25, 1.25, 1.4, 1.3, 1.3, 1.3, 1.4, 1.2, 1.35, 1.35, 1.35, 1.35, 1.25],
    )
    r0, r1, r2, r3 = [table.rows[i].cells for i in range(4)]

    r0[0].text = "№ потребителя"
    r0[1].text = "Наименование потребителя"
    r0[2].text = "Техпроцесс"
    r0[3].text = "Время, ч"
    r0[4].text = "Количество"
    r0[5].text = "Обоснование"
    r0[6].text = "Расход на единицу, м³/сут"
    r0[7].text = "Качество воды"
    r0[8].text = "Общее водопотребление, м³/сут"
    r0[9].text = "Водопотребление, м³/сут"
    r0[13].text = "Безвозвратные потери,\nм³/сут"
    r0[14].text = "Водоотведение,\nм³/сут"

    r1[9].text = "Источники водоснабжения"
    r1[14].text = "Городская канализация"
    r1[18].text = "Водосток"

    r2[9].text = "Горводопровод"
    r2[10].text = "Скважины"
    r2[11].text = "Техвода"
    r2[12].text = "Оборотные системы"
    r2[14].text = "Хоз.-быт."
    r2[15].text = "Норм.-чистые"
    r2[16].text = "Загр. мех./мин."
    r2[17].text = "Загр. хим./орг."

    for c in range(0, 9):
        table.cell(0, c).merge(table.cell(2, c))
    table.cell(0, 9).merge(table.cell(0, 12))
    table.cell(1, 9).merge(table.cell(1, 12))
    table.cell(0, 13).merge(table.cell(2, 13))
    table.cell(0, 14).merge(table.cell(0, 18))
    table.cell(1, 14).merge(table.cell(1, 17))
    table.cell(1, 18).merge(table.cell(2, 18))

    for i in range(19):
        r3[i].text = str(i + 1)
        r3[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for rr in range(0, 4):
        for cc in range(0, 19):
            _set_cell_center(table.cell(rr, cc))

    # Вертикальные подписи по форме: столбцы 2-9, 14 и 19.
    # "Водоотведение" (столбец 15) оставляем горизонтально.
    for cidx in [1, 2, 3, 4, 5, 6, 7, 8, 13]:
        _set_cell_vertical(table.cell(0, cidx))
        _set_cell_center(table.cell(0, cidx))

    # Вертикальные слова из подзаголовков.
    for cidx in [9, 10, 11, 12, 14, 15, 16, 17]:
        _set_cell_vertical(table.cell(2, cidx))
        _set_cell_center(table.cell(2, cidx))
    _set_cell_vertical(table.cell(2, 18))
    _set_cell_center(table.cell(2, 18))

    np_source = str(water_options.get("np_water_source", "Горводопровод"))
    np_storm = _optf(water_options, "np_storm_m3_day", 0.0)

    source_col_default = {
        "Горводопровод": 9,
        "Скважины": 10,
        "Техвода": 11,
        "Оборотные системы": 12,
    }.get(np_source, 9)
    sewer_col_map = {
        "Хоз.-быт.": 14,
        "Норм.-чистые": 15,
        "Загр. мех./мин.": 16,
        "Загр. хим./орг.": 17,
    }
    sum_source_9 = 0.0
    sum_source_10 = 0.0
    sum_source_11 = 0.0
    sum_source_12 = 0.0
    sum_loss = 0.0
    sum_sewer_14 = 0.0
    sum_sewer_15 = 0.0
    sum_sewer_16 = 0.0
    sum_sewer_17 = 0.0

    for idx, c in enumerate(consumers, start=1):
        if float(c.get("count", 0.0) or 0.0) <= 0:
            continue
        row = table.add_row().cells
        name = str(c.get("name", ""))
        q_total = float(c.get("total_m3_day", 0.0) or 0.0)
        norm = (float(c.get("cold_l_per_unit_day", 0.0) or 0.0) + float(c.get("hot_l_per_unit_day", 0.0) or 0.0)) / 1000.0
        is_irrig = _is_irrigation_name(name)
        source_name = str(c.get("np_source_override", "") or "").strip()
        source_col = {
            "Горводопровод": 9,
            "Скважины": 10,
            "Техвода": 11,
            "Оборотные системы": 12,
        }.get(source_name, source_col_default)
        sewer_override = str(c.get("np_sewer_override", "") or "").strip()
        sewer_col = sewer_col_map.get(sewer_override, _infer_form2_sewer_col(name))
        water_quality = str(c.get("water_quality_override", "") or "").strip()
        if not water_quality:
            water_quality = "Питьевая"
        row[0].text = str(idx)
        row[1].text = name
        row[2].text = ""
        name = str(c.get("name", ""))
        row[3].text = "" if _is_no_time_consumer(name) else _fmt_or_blank(float(c.get("t_hours", 24.0) or 24.0))
        row[4].text = _fmt2(float(c.get("count", 0.0) or 0.0))
        row[5].text = ""
        row[6].text = _fmt_or_blank(norm)
        row[7].text = water_quality
        row[8].text = _fmt2(q_total)
        row[9].text = _fmt_or_blank(q_total if source_col == 9 else 0.0)
        row[10].text = _fmt_or_blank(q_total if source_col == 10 else 0.0)
        row[11].text = _fmt_or_blank(q_total if source_col == 11 else 0.0)
        row[12].text = _fmt_or_blank(q_total if source_col == 12 else 0.0)
        row[13].text = _fmt_or_blank(q_total if is_irrig else 0.0)
        sewer_q = 0.0 if is_irrig else q_total
        row[14].text = _fmt_or_blank(sewer_q if sewer_col == 14 else 0.0)
        row[15].text = _fmt_or_blank(sewer_q if sewer_col == 15 else 0.0)
        row[16].text = _fmt_or_blank(sewer_q if sewer_col == 16 else 0.0)
        row[17].text = _fmt_or_blank(sewer_q if sewer_col == 17 else 0.0)
        row[18].text = ""

        if source_col == 9:
            sum_source_9 += q_total
        elif source_col == 10:
            sum_source_10 += q_total
        elif source_col == 11:
            sum_source_11 += q_total
        else:
            sum_source_12 += q_total

        sum_loss += q_total if is_irrig else 0.0
        if not is_irrig:
            if sewer_col == 14:
                sum_sewer_14 += q_total
            elif sewer_col == 15:
                sum_sewer_15 += q_total
            elif sewer_col == 16:
                sum_sewer_16 += q_total
            else:
                sum_sewer_17 += q_total

    total = table.add_row().cells
    q_day = float(water_results.get("total_m3_day", 0.0))
    total[1].text = "ИТОГ"
    total[8].text = _fmt2(q_day)
    total[9].text = _fmt_or_blank(sum_source_9)
    total[10].text = _fmt_or_blank(sum_source_10)
    total[11].text = _fmt_or_blank(sum_source_11)
    total[12].text = _fmt_or_blank(sum_source_12)
    total[13].text = _fmt_or_blank(sum_loss)
    total[14].text = _fmt_or_blank(sum_sewer_14)
    total[15].text = _fmt_or_blank(sum_sewer_15)
    total[16].text = _fmt_or_blank(sum_sewer_16)
    total[17].text = _fmt_or_blank(sum_sewer_17)
    total[18].text = _fmt_or_blank(np_storm)

    _set_table_font_size(table, 11)


def _add_form1_balance_table(
    doc: Document,
    object_name: str,
    object_address: str,
    consumers: List[Dict[str, float | str]],
    water_results: Dict[str, float],
    water_options: Dict[str, str],
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("БАЛАНС ВОДОСНАБЖЕНИЯ И ВОДООТВЕДЕНИЯ").bold = True
    obj_line = object_name.strip() if object_name else "______________________________"
    addr_line = object_address.strip() if object_address else "______________________________"
    doc.add_paragraph(f"Наименование объекта: {obj_line}")
    doc.add_paragraph(f"Адрес объекта: {addr_line}")

    table = doc.add_table(rows=4, cols=24)
    table.style = "Table Grid"
    _set_col_widths(
        table,
        [0.5, 2.5, 1.1, 0.9, 1.5, 1.2, 1.2, 1.0, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95, 0.9, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95, 0.8, 1.2, 1.0],
    )
    r0, r1, r2, r3 = [table.rows[i].cells for i in range(4)]

    r0[0].text = "№"
    r0[1].text = "Наименование потребителя"
    r0[2].text = "Количество потребителей"
    r0[3].text = "Время работы в сутки, ч"
    r0[4].text = "Требования к качеству воды"
    r0[5].text = "Необходимый напор Hтр (п. 8.27), м"
    r0[6].text = "Режим водопотребления"
    r0[7].text = "Расход воды на одного потребителя,\nм³/ч"
    r0[8].text = "Водопотребление"
    r0[14].text = "Характеристика сточных вод"
    r0[15].text = "Режим водоотведения"
    r0[16].text = "Водоотведение, м³/сут"
    r0[22].text = "Концентрация загрязнений локальных очистных сооружений, мг/л"
    r0[23].text = "Примечание"

    r1[8].text = "из хозяйственно-\nпитьевого\nводопровода"
    r1[11].text = "из производственного\nводопровода"
    r1[16].text = "в бытовую\nканализацию"
    r1[19].text = "в производственную\nканализацию"

    r2[8].text = "м³/сут"
    r2[9].text = "м³/ч"
    r2[10].text = "л/с"
    r2[11].text = "м³/сут"
    r2[12].text = "м³/ч"
    r2[13].text = "л/с"
    r2[14].text = ""
    r2[16].text = "м³/сут"
    r2[17].text = "м³/ч"
    r2[18].text = "л/с"
    r2[19].text = "м³/сут"
    r2[20].text = "м³/ч"
    r2[21].text = "л/с"

    for c in range(0, 8):
        table.cell(0, c).merge(table.cell(2, c))
    # Водопотребление: только графы 9-14 (колонки 8-13)
    table.cell(0, 8).merge(table.cell(0, 13))
    table.cell(1, 8).merge(table.cell(1, 10))
    table.cell(1, 11).merge(table.cell(1, 13))
    table.cell(0, 14).merge(table.cell(2, 14))
    table.cell(0, 15).merge(table.cell(2, 15))
    table.cell(0, 16).merge(table.cell(0, 21))
    table.cell(1, 16).merge(table.cell(1, 18))
    table.cell(1, 19).merge(table.cell(1, 21))
    table.cell(0, 22).merge(table.cell(2, 22))
    table.cell(0, 23).merge(table.cell(2, 23))

    for i in range(24):
        r3[i].text = str(i + 1)
        r3[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for rr in range(0, 4):
        for cc in range(0, 24):
            _set_cell_center(table.cell(rr, cc))

    # Явная схема ориентации как в форме ГОСТ (без глобального разворота всей шапки).
    # Вертикально: одиночные столбцы 1-8, 15,16,23,24.
    for cidx in [0, 1, 2, 3, 4, 5, 6, 7, 14, 15, 22, 23]:
        _set_cell_vertical(table.cell(0, cidx))
        _set_cell_center(table.cell(0, cidx))

    # Горизонтально: групповые заголовки.
    _set_cell_horizontal(table.cell(0, 8))   # Водопотребление
    _set_cell_horizontal(table.cell(0, 16))  # Водоотведение

    # Горизонтально: 4 подзаголовка источников/канализации.
    for cidx in [8, 11, 16, 19]:
        _set_cell_horizontal(table.cell(1, cidx))

    # Единицы измерения в строке 3 — горизонтально, как в форме.
    for cidx in [8, 9, 10, 11, 12, 13, 16, 17, 18, 19, 20, 21]:
        _set_cell_horizontal(table.cell(2, cidx))

    pr_source = str(water_options.get("pr_water_source", "Из хоз.-питьевого водопровода"))
    pr_conc = str(water_options.get("pr_concentration_mg_l", "")).strip()
    is_prod_source = pr_source == "Из производственного водопровода"
    # Графа 6 формы 1: приоритет у расчетного требуемого напора Hтр (п. 8.27),
    # fallback — ручной ввод давления на вводе.
    hyd_required_head_m = _optf(water_options, "hyd_required_head_m_hvs", 0.0)
    pr_inlet_pressure = str(water_options.get("pr_inlet_pressure_mpa", "") or "").strip()
    pr_inlet_pressure = "" if pr_inlet_pressure in ("0", "0.0", "0.00", "0.000") else pr_inlet_pressure
    col6_value = _fmt_or_blank_prec(hyd_required_head_m, 2) if hyd_required_head_m > 0.0 else pr_inlet_pressure
    for idx, c in enumerate(consumers, start=1):
        if float(c.get("count", 0.0) or 0.0) <= 0:
            continue
        q_total = float(c.get("total_m3_day", 0.0) or 0.0)
        q_max_h = float(c.get("total_max_m3_hour", 0.0) or 0.0)
        if q_max_h <= 0.0:
            q_max_h = q_total * 1.8 / 24.0
        q_sec = q_max_h * 1000.0 / 3600.0
        q_u_total_l_day = float(c.get("q_u_total_l_day", 0.0) or 0.0)
        q_one_m3_h = q_u_total_l_day / 1000.0 / 24.0 if q_u_total_l_day > 0 else 0.0
        row = table.add_row().cells
        row[0].text = str(idx)
        name = str(c.get("name", ""))
        row[1].text = name
        row[2].text = _fmt2(float(c.get("count", 0.0) or 0.0))
        t_hours = float(c.get("t_hours", 24.0) or 24.0)
        row[3].text = "" if _is_no_time_consumer(name) else _fmt2(t_hours)
        row[4].text = _infer_prod_water_quality(
            name,
            str(c.get("water_quality_override", "") or "").strip(),
        )
        row[5].text = col6_value if col6_value else ""
        row[6].text = _infer_consumption_mode(name, t_hours)
        row[7].text = _fmt_or_blank_prec(q_one_m3_h, 5)
        row_is_prod_source = is_prod_source or bool(c.get("use_prod_water_source", False))
        row[8].text = _fmt_or_blank(0.0 if row_is_prod_source else q_total)
        row[9].text = _fmt_or_blank(0.0 if row_is_prod_source else q_max_h)
        row[10].text = _fmt_or_blank(0.0 if row_is_prod_source else q_sec)
        row[11].text = _fmt_or_blank(q_total if row_is_prod_source else 0.0)
        row[12].text = _fmt_or_blank(q_max_h if row_is_prod_source else 0.0)
        row[13].text = _fmt_or_blank(q_sec if row_is_prod_source else 0.0)
        sewer_target = str(c.get("sewer_target", "domestic")).strip().lower()
        is_prod_sewer = sewer_target == "production"
        water_quality = _infer_prod_water_quality(
            name,
            str(c.get("water_quality_override", "") or "").strip(),
        )
        row[14].text = _infer_sewer_characteristic(name, sewer_target, water_quality)
        row[15].text = _infer_consumption_mode(name, t_hours)
        row[16].text = _fmt_or_blank(0.0 if is_prod_sewer else q_total)
        row[17].text = _fmt_or_blank(0.0 if is_prod_sewer else q_max_h)
        row[18].text = _fmt_or_blank(0.0 if is_prod_sewer else q_sec)
        row[19].text = _fmt_or_blank(q_total if is_prod_sewer else 0.0)
        row[20].text = _fmt_or_blank(q_max_h if is_prod_sewer else 0.0)
        row[21].text = _fmt_or_blank(q_sec if is_prod_sewer else 0.0)
        row[22].text = pr_conc
        row[23].text = "безвозвратные потери" if _is_irrigation_name(name) else ""

    def _sum_sewer(component: str, route: str) -> float:
        return sum(
            float(c.get(component, 0.0) or 0.0)
            for c in consumers
            if float(c.get("count", 0.0) or 0.0) > 0 and str(c.get("sewer_target", "domestic")).strip().lower() == route
        )

    sewer_dom_day_c = _sum_sewer("cold_m3_day", "domestic")
    sewer_dom_day_h = _sum_sewer("hot_m3_day", "domestic")
    sewer_dom_day_t = _sum_sewer("total_m3_day", "domestic")
    sewer_dom_h_c = _sum_sewer("cold_max_m3_hour", "domestic")
    sewer_dom_h_h = _sum_sewer("hot_max_m3_hour", "domestic")
    sewer_dom_h_t = _sum_sewer("total_max_m3_hour", "domestic")
    sewer_dom_s_c = sewer_dom_h_c * 1000.0 / 3600.0
    sewer_dom_s_h = sewer_dom_h_h * 1000.0 / 3600.0
    sewer_dom_s_t = sewer_dom_h_t * 1000.0 / 3600.0

    sewer_pr_day_c = _sum_sewer("cold_m3_day", "production")
    sewer_pr_day_h = _sum_sewer("hot_m3_day", "production")
    sewer_pr_day_t = _sum_sewer("total_m3_day", "production")
    sewer_pr_h_c = _sum_sewer("cold_max_m3_hour", "production")
    sewer_pr_h_h = _sum_sewer("hot_max_m3_hour", "production")
    sewer_pr_h_t = _sum_sewer("total_max_m3_hour", "production")
    sewer_pr_s_c = sewer_pr_h_c * 1000.0 / 3600.0
    sewer_pr_s_h = sewer_pr_h_h * 1000.0 / 3600.0
    sewer_pr_s_t = sewer_pr_h_t * 1000.0 / 3600.0

    total = table.add_row().cells
    qd_hh = 0.0
    qh_hh = 0.0
    qs_hh = 0.0
    qd_pr = 0.0
    qh_pr = 0.0
    qs_pr = 0.0
    for c in consumers:
        if float(c.get("count", 0.0) or 0.0) <= 0:
            continue
        qd_i = float(c.get("total_m3_day", 0.0) or 0.0)
        qh_i = float(c.get("total_max_m3_hour", 0.0) or 0.0)
        if qh_i <= 0.0:
            qh_i = qd_i * 1.8 / 24.0
        qs_i = float(c.get("total_max_l_sec", 0.0) or 0.0)
        if qs_i <= 0.0:
            qs_i = qh_i * 1000.0 / 3600.0
        row_is_prod_source = is_prod_source or bool(c.get("use_prod_water_source", False))
        if row_is_prod_source:
            qd_pr += qd_i
            qh_pr += qh_i
            qs_pr += qs_i
        else:
            qd_hh += qd_i
            qh_hh += qh_i
            qs_hh += qs_i
    total[1].text = "ИТОГ"
    total[8].text = _fmt_or_blank(qd_hh)
    total[9].text = _fmt_or_blank(qh_hh)
    total[10].text = _fmt_or_blank(qs_hh)
    total[11].text = _fmt_or_blank(qd_pr)
    total[12].text = _fmt_or_blank(qh_pr)
    total[13].text = _fmt_or_blank(qs_pr)
    total[15].text = ""
    total[16].text = _fmt_or_blank(sewer_dom_day_t)
    total[17].text = _fmt_or_blank(sewer_dom_h_t)
    total[18].text = _fmt_or_blank(sewer_dom_s_t)
    total[19].text = _fmt_or_blank(sewer_pr_day_t)
    total[20].text = _fmt_or_blank(sewer_pr_h_t)
    total[21].text = _fmt_or_blank(sewer_pr_s_t)
    total[22].text = pr_conc

    _set_table_font_size(table, 11)


def _add_checks_block(doc: Document, checks: List[str]) -> None:
    if not checks:
        return
    doc.add_paragraph("Контроль заполнения:")
    for check in checks:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(check)


def _add_detailed_water_calc_table(doc: Document, consumers: List[Dict[str, float | str]], water_results: Dict[str, float]) -> None:
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr.add_run("Расчет водоснабжения и водоотведения").bold = True
    doc.add_paragraph("Расчет выполнен в соответствии с СП 30.13330.2020.")

    active = [c for c in consumers if float(c.get("count", 0.0) or 0.0) > 0.0]
    table = doc.add_table(rows=4, cols=15)
    table.style = "Table Grid"

    h0 = table.rows[0].cells
    h1 = table.rows[1].cells
    h2 = table.rows[2].cells
    h3 = table.rows[3].cells

    h0[0].text = "Наименование\nводопотребителей"
    h0[1].text = "коли-\nчество\nU\nсутки\nчас"
    h0[2].text = "нормы рас-\nхода воды"
    h0[4].text = "расход воды\nприбором"
    h0[6].text = "расход воды\nводопотребителями"
    h0[9].text = "NP"
    h0[10].text = "NPhr"
    h0[11].text = "α"
    h0[12].text = "αhr"
    h0[13].text = "макси-\nмальный\nрасчетный\nрасход\n5 · q0 · α"
    h0[14].text = "макси-\nмальный\nчасовой\nрасход\n0.005 · q0,hr · αhr"

    h1[2].text = "сутки"
    h1[3].text = "час"
    h1[4].text = "час"
    h1[5].text = "сек"
    h1[6].text = "сутки"
    h1[7].text = "час"
    h1[8].text = "ср.час"

    h2[2].text = "qᶜu\nqʰu\nл/сут"
    h2[3].text = "qᶜhr,u\nqʰhr,u\nл/ч"
    h2[4].text = "qᶜ0,hr\nqʰ0,hr\nл/ч"
    h2[5].text = "qᶜ0\nqʰ0\nл/с"
    h2[6].text = "qᶜ0 · U / 1000\nqʰu · U / 1000\nм³/сут"
    h2[7].text = "qᶜhr · U\nqʰhr · U\nл/ч"
    h2[8].text = "qᶜT\nqʰT\nм³/ч"
    h2[9].text = "qhr,u · U\nq0 · 3600"
    h2[10].text = "qhr,u · U\nq0,hr"
    h2[13].text = "qᶜ, qʰ\nл/с"
    h2[14].text = "qᶜhr, qʰhr\nм³/ч"

    for i in range(15):
        h3[i].text = str(i + 1)
        _set_cell_center(h3[i])

    table.cell(0, 0).merge(table.cell(2, 0))
    table.cell(0, 1).merge(table.cell(2, 1))
    table.cell(0, 2).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(0, 5))
    table.cell(0, 6).merge(table.cell(0, 8))
    for col in (9, 10, 11, 12, 13, 14):
        table.cell(0, col).merge(table.cell(2, col))

    for rr in range(0, 4):
        for cc in range(0, 15):
            _set_cell_center(table.cell(rr, cc))

    def _fmt_local(v: float, decimals: int = 2) -> str:
        return f"{float(v):.{decimals}f}".replace(".", ",")

    def _fmt_dash(v: float, decimals: int = 2) -> str:
        if abs(float(v or 0.0)) < 1e-9:
            return "-"
        return _fmt_local(v, decimals)

    def _section_values(c: Dict[str, float | str], sec: str) -> Dict[str, float]:
        u = float(c.get("count", 0.0) or 0.0)
        t_h = max(float(c.get("t_hours", 24.0) or 24.0), 0.01)
        if sec == "cold":
            q_u = float(c.get("cold_l_per_unit_day", 0.0) or 0.0)
            q_hr = max(float(c.get("q_hr_total_l_h", 0.0) or 0.0) - float(c.get("q_hr_hot_l_h", 0.0) or 0.0), 0.0)
            q0hr = float(c.get("q0_spec_l_h", 0.0) or 0.0)
            q0 = float(c.get("q0_spec_l_s", 0.0) or 0.0)
            q_day = float(c.get("cold_m3_day", 0.0) or 0.0)
            alpha = float(c.get("alpha_cold", 0.0) or 0.0)
            alpha_hr = float(c.get("alpha_hr_cold", 0.0) or 0.0)
            qmax_s = float(c.get("cold_max_l_sec", 0.0) or 0.0)
            qmax_h = float(c.get("cold_max_m3_hour", 0.0) or 0.0)
            np = float(c.get("np_cold", 0.0) or 0.0)
            nphr = float(c.get("np_hr_cold", 0.0) or 0.0)
        elif sec == "hot":
            q_u = float(c.get("hot_l_per_unit_day", 0.0) or 0.0)
            q_hr = float(c.get("q_hr_hot_l_h", 0.0) or 0.0)
            q0hr = float(c.get("q0_spec_l_h", 0.0) or 0.0)
            q0 = float(c.get("q0_spec_l_s", 0.0) or 0.0)
            q_day = float(c.get("hot_m3_day", 0.0) or 0.0)
            alpha = float(c.get("alpha_hot", 0.0) or 0.0)
            alpha_hr = float(c.get("alpha_hr_hot", 0.0) or 0.0)
            qmax_s = float(c.get("hot_max_l_sec", 0.0) or 0.0)
            qmax_h = float(c.get("hot_max_m3_hour", 0.0) or 0.0)
            np = float(c.get("np_hot", 0.0) or 0.0)
            nphr = float(c.get("np_hr_hot", 0.0) or 0.0)
        else:
            q_u = float(c.get("q_u_total_l_day", 0.0) or 0.0)
            q_hr = float(c.get("q_hr_total_l_h", 0.0) or 0.0)
            q0hr = float(c.get("q0_total_l_h", 0.0) or 0.0)
            q0 = float(c.get("q0_total_l_s", 0.0) or 0.0)
            q_day = float(c.get("total_m3_day", 0.0) or 0.0)
            alpha = float(c.get("alpha_total", 0.0) or 0.0)
            alpha_hr = float(c.get("alpha_hr_total", 0.0) or 0.0)
            qmax_s = float(c.get("total_max_l_sec", 0.0) or 0.0)
            qmax_h = float(c.get("total_max_m3_hour", 0.0) or 0.0)
            np = float(c.get("np_total", 0.0) or 0.0)
            nphr = float(c.get("np_hr_total", 0.0) or 0.0)
        q_hr_u = q_hr * u
        q_avg_h = q_day / t_h
        return {
            "u": u,
            "q_u": q_u,
            "q_hr": q_hr,
            "q0hr": q0hr,
            "q0": q0,
            "q_day": q_day,
            "q_hr_u": q_hr_u,
            "q_avg_h": q_avg_h,
            "np": np,
            "nphr": nphr,
            "alpha": alpha,
            "alpha_hr": alpha_hr,
            "qmax_s": qmax_s,
            "qmax_h": qmax_h,
        }

    def _special_kind(name: str) -> str:
        n = (name or "").strip().lower()
        if "подпитка котельной" in n:
            return "boiler"
        if "полив" in n or "заливка поверхности катка" in n:
            return "irrigation"
        return ""

    def _add_section(title: str, sec: str) -> None:
        cap = table.add_row().cells
        for i in range(1, 15):
            cap[i].text = ""
        merged_cap = cap[0].merge(cap[14])
        _set_cell_text_center(merged_cap, title, bold=True)

        sum_day = 0.0
        sum_hr_u = 0.0
        sum_avg_h = 0.0
        sum_np = 0.0
        sum_nphr = 0.0
        sum_alpha = 0.0
        sum_alpha_hr = 0.0
        sum_qmax_s = 0.0
        sum_qmax_h = 0.0
        nrows = 0
        first_q0 = 0.0
        first_q0hr = 0.0

        core_rows = [c for c in active if _special_kind(str(c.get("name", ""))) == ""]
        special_rows = [c for c in active if _special_kind(str(c.get("name", ""))) != ""]

        for c in core_rows:
            v = _section_values(c, sec)
            row = table.add_row().cells
            row[0].text = str(c.get("name", ""))
            row[1].text = _fmt_local(v["u"], 0)
            row[2].text = _fmt_local(v["q_u"], 1)
            row[3].text = _fmt_local(v["q_hr"], 1)
            row[4].text = _fmt_local(v["q0hr"], 0)
            row[5].text = _fmt_local(v["q0"], 2)
            row[6].text = _fmt_local(v["q_day"], 2)
            row[7].text = _fmt_local(v["q_hr_u"], 0)
            row[8].text = _fmt_local(v["q_avg_h"], 2)
            row[9].text = _fmt_local(v["np"], 2)
            row[10].text = _fmt_local(v["nphr"], 1)
            row[11].text = _fmt_local(v["alpha"], 3)
            row[12].text = _fmt_local(v["alpha_hr"], 3)
            row[13].text = _fmt_local(v["qmax_s"], 2)
            row[14].text = _fmt_local(v["qmax_h"], 2)

            if first_q0 <= 0.0:
                first_q0 = v["q0"]
            if first_q0hr <= 0.0:
                first_q0hr = v["q0hr"]
            sum_day += v["q_day"]
            sum_hr_u += v["q_hr_u"]
            sum_avg_h += v["q_avg_h"]
            sum_np += v["np"]
            sum_nphr += v["nphr"]
            sum_alpha += v["alpha"]
            sum_alpha_hr += v["alpha_hr"]
            sum_qmax_s += v["qmax_s"]
            sum_qmax_h += v["qmax_h"]
            nrows += 1

        qrow = table.add_row().cells
        for i in range(13):
            qrow[i].text = ""
        qrow[13].text = f"q0={_fmt_local(first_q0, 2)}" if first_q0 > 0 else ""
        qrow[14].text = f"q0hr={_fmt_local(first_q0hr, 0)}" if first_q0hr > 0 else ""

        subtotal = table.add_row().cells
        merged_subtotal = subtotal[0].merge(subtotal[5])
        _set_cell_text_center(merged_subtotal, "Итог - хозяйственно-питьевые нужды:")
        subtotal[6].text = _fmt_local(sum_day, 2)
        subtotal[7].text = _fmt_local(sum_hr_u, 0)
        subtotal[8].text = _fmt_local(sum_avg_h, 2)
        subtotal[9].text = _fmt_local(sum_np, 2)
        subtotal[10].text = _fmt_local(sum_nphr, 1)
        subtotal[11].text = _fmt_local(sum_alpha / max(nrows, 1), 3)
        subtotal[12].text = _fmt_local(sum_alpha_hr / max(nrows, 1), 3)
        subtotal[13].text = _fmt_local(sum_qmax_s, 2)
        subtotal[14].text = _fmt_local(sum_qmax_h, 2)

        add_day = 0.0
        add_avg_h = 0.0
        add_qmax_s = 0.0
        add_qmax_h = 0.0

        for c in special_rows:
            kind = _special_kind(str(c.get("name", "")))
            v = _section_values(c, sec)

            # Полив показываем отдельно и не смешиваем с общим "total" разделом.
            if kind == "irrigation" and sec == "total":
                row = table.add_row().cells
                row[0].text = str(c.get("name", ""))
                row[1].text = _fmt_local(v["u"], 0)
                for i in range(2, 15):
                    row[i].text = "-"
                continue

            row = table.add_row().cells
            row[0].text = str(c.get("name", ""))
            row[1].text = _fmt_local(v["u"], 0)
            row[2].text = _fmt_dash(v["q_u"], 1)
            row[3].text = _fmt_dash(v["q_hr"], 1)
            row[4].text = _fmt_dash(v["q0hr"], 0)
            row[5].text = _fmt_dash(v["q0"], 2)
            row[6].text = _fmt_dash(v["q_day"], 2)
            row[7].text = _fmt_dash(v["q_hr_u"], 0)
            row[8].text = _fmt_dash(v["q_avg_h"], 2)
            row[9].text = "-"
            row[10].text = "-"
            row[11].text = "-"
            row[12].text = "-"
            row[13].text = _fmt_dash(v["qmax_s"], 2)
            row[14].text = _fmt_dash(v["qmax_h"], 2)

            add_day += v["q_day"]
            add_avg_h += v["q_avg_h"]
            add_qmax_s += v["qmax_s"]
            add_qmax_h += v["qmax_h"]

        total = table.add_row().cells
        merged_total = total[0].merge(total[5])
        _set_cell_text_center(merged_total, "Итог:")
        total[6].text = _fmt_local(sum_day + add_day, 2)
        total[7].text = "-"
        total[8].text = _fmt_local(sum_avg_h + add_avg_h, 2)
        total[9].text = "-"
        total[10].text = "-"
        total[11].text = "-"
        total[12].text = "-"
        total[13].text = _fmt_local(sum_qmax_s + add_qmax_s, 2)
        total[14].text = _fmt_local(sum_qmax_h + add_qmax_h, 2)

    _add_section("Расчет расходов холодной воды", "cold")
    _add_section("Расчет расходов горячей воды", "hot")
    _add_section("Расчет расходов воды общей", "total")

    _set_table_font_size(table, 11)


def build_report_docx(
    project_name: str,
    object_name: str,
    object_address: str,
    annex_label: str,
    project_meta: Dict[str, str],
    water_inputs: Dict[str, str],
    water_results: Dict[str, float],
    water_consumers: List[Dict[str, float | str]],
    gvs_results: Dict[str, float],
    checks: List[str],
) -> bytes:
    doc = Document()
    _set_doc_defaults(doc)

    annex_text = (annex_label or "").strip()
    p_annex = doc.add_paragraph()
    p_annex.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_annex.add_run(f"Приложение {annex_text or '____'}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Расчет выполнен на основе СП 30.13330.2020 (приложение А.2), с учетом действующих изменений.").bold = False
    doc.add_paragraph("Таблица оформлена по ГОСТ Р 21.619-2023, приложение А.")
    doc.add_paragraph()

    if _is_production_object(water_consumers):
        _add_form1_balance_table(
            doc,
            object_name=object_name,
            object_address=object_address,
            consumers=water_consumers,
            water_results=water_results,
            water_options=water_inputs,
        )
    else:
        _add_form2_balance_table(
            doc,
            object_name=object_name,
            object_address=object_address,
            consumers=water_consumers,
            water_results=water_results,
            water_options=water_inputs,
        )
    doc.add_page_break()
    _add_detailed_water_calc_table(doc, consumers=water_consumers, water_results=water_results)

    doc.add_paragraph()
    _add_kv_table(
        doc,
        [
            ("QTh, средний тепловой поток ГВС, кВт", f'{gvs_results.get("qth_kW", 0.0):.3f}'),
            ("Qhr,h, максимальный тепловой поток ГВС, кВт", f'{gvs_results.get("qhrh_kW", 0.0):.3f}'),
            ("qcir, расход циркуляции, л/с", f'{gvs_results.get("qcir_l_s", 0.0):.4f}'),
            ("qh,cir, расход с циркуляцией, л/с", f'{gvs_results.get("qh_cir_l_s", 0.0):.4f}'),
        ],
    )
    _add_checks_block(doc, checks)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
